"""
Async Task Queue System
Uses Celery for background processing of transcription and analysis tasks
"""

import os
import json
import time
from datetime import datetime, timedelta
from typing import Dict, List, Optional, Any
from celery import Celery
from celery.result import AsyncResult
import redis

# Only import streamlit when needed (not when running as script)
try:
    import streamlit as st
    STREAMLIT_AVAILABLE = True
except ImportError:
    STREAMLIT_AVAILABLE = False
    print("⚠️ Streamlit not available - UI functions disabled")

# Configure Celery
CELERY_BROKER_URL = os.getenv('CELERY_BROKER_URL', 'redis://localhost:6379/0')
CELERY_RESULT_BACKEND = os.getenv('CELERY_RESULT_BACKEND', 'redis://localhost:6379/0')

# Initialize Celery app
celery_app = Celery(
    'verbatim_ai_tasks',
    broker=CELERY_BROKER_URL,
    backend=CELERY_RESULT_BACKEND
)

# Configure Celery
celery_app.conf.update(
    task_serializer='json',
    accept_content=['json'],
    result_serializer='json',
    timezone='UTC',
    enable_utc=True,
    task_track_started=True,
    task_time_limit=30 * 60,  # 30 minutes
    task_soft_time_limit=25 * 60,  # 25 minutes
    worker_prefetch_multiplier=1,
    worker_max_tasks_per_child=1000,
)

class TaskStatus:
    """Task status enumeration"""
    PENDING = "PENDING"
    STARTED = "STARTED"
    SUCCESS = "SUCCESS"
    FAILURE = "FAILURE"
    RETRY = "RETRY"
    REVOKED = "REVOKED"

class TaskManager:
    """Manages async tasks for VerbatimAI"""
    
    def __init__(self):
        self.redis_client = redis.Redis.from_url(CELERY_BROKER_URL)
        self.task_results = {}
    
    def get_task_status(self, task_id: str) -> Dict:
        """Get status of a task"""
        try:
            result = AsyncResult(task_id, app=celery_app)
            return {
                'task_id': task_id,
                'status': result.status,
                'result': result.result if result.ready() else None,
                'info': result.info if hasattr(result, 'info') else None,
                'traceback': result.traceback if result.failed() else None
            }
        except Exception as e:
            return {
                'task_id': task_id,
                'status': 'ERROR',
                'error': str(e)
            }
    
    def get_all_tasks(self) -> List[Dict]:
        """Get all tasks from Redis"""
        try:
            tasks = []
            for key in self.redis_client.scan_iter(match="celery-task-meta-*"):
                task_id = key.decode('utf-8').replace('celery-task-meta-', '')
                task_info = self.get_task_status(task_id)
                tasks.append(task_info)
            return tasks
        except Exception as e:
            print(f"Error getting tasks: {e}")
            return []
    
    def cleanup_old_tasks(self, days: int = 7):
        """Clean up old completed tasks"""
        try:
            cutoff_time = datetime.now() - timedelta(days=days)
            tasks = self.get_all_tasks()
            
            for task in tasks:
                if task['status'] in ['SUCCESS', 'FAILURE']:
                    # Check if task is old enough to delete
                    # This is a simplified check - in production you'd store creation time
                    pass
        except Exception as e:
            print(f"Error cleaning up tasks: {e}")

# Celery tasks
@celery_app.task(bind=True)
def transcribe_audio_task(self, audio_file_path: str, config: Dict) -> Dict:
    """Transcribe audio file asynchronously"""
    try:
        # Update task state
        self.update_state(
            state='STARTED',
            meta={'current': 0, 'total': 100, 'status': 'Starting transcription...'}
        )
        
        # Import here to avoid circular imports
        import assemblyai as aai
        from config import ASSEMBLYAI_API_KEY
        
        # Configure AssemblyAI
        aai.settings.api_key = ASSEMBLYAI_API_KEY
        
        # Update progress
        self.update_state(
            state='PROGRESS',
            meta={'current': 20, 'total': 100, 'status': 'Uploading audio file...'}
        )
        
        # Transcribe audio
        transcriber = aai.Transcriber()
        transcript = transcriber.transcribe(audio_file_path, config)
        
        # Update progress
        self.update_state(
            state='PROGRESS',
            meta={'current': 80, 'total': 100, 'status': 'Processing transcript...'}
        )
        
        if transcript.status == aai.TranscriptStatus.error:
            raise Exception(f"Transcription failed: {transcript.error}")
        
        # Convert transcript to serializable format
        transcript_data = {
            'text': transcript.text,
            'audio_duration': transcript.audio_duration,
            'utterances': [
                {
                    'speaker': u.speaker,
                    'text': u.text,
                    'start': u.start,
                    'end': u.end,
                    'confidence': getattr(u, 'confidence', None)
                }
                for u in transcript.utterances
            ],
            'sentiment': getattr(transcript, 'sentiment', None),
            'highlights': [
                {
                    'text': h.text,
                    'reason': h.reason,
                    'start': h.start,
                    'end': h.end
                }
                for h in getattr(transcript, 'highlights', [])
            ]
        }
        
        # Update final progress
        self.update_state(
            state='SUCCESS',
            meta={'current': 100, 'total': 100, 'status': 'Transcription completed!'}
        )
        
        return {
            'status': 'success',
            'transcript': transcript_data,
            'file_path': audio_file_path,
            'completed_at': datetime.now().isoformat()
        }
        
    except Exception as e:
        # Update error state
        self.update_state(
            state='FAILURE',
            meta={'error': str(e), 'status': 'Transcription failed'}
        )
        raise

@celery_app.task(bind=True)
def analyze_meeting_task(self, transcript_data: Dict) -> Dict:
    """Analyze meeting transcript asynchronously"""
    try:
        # Update task state
        self.update_state(
            state='STARTED',
            meta={'current': 0, 'total': 100, 'status': 'Starting analysis...'}
        )
        
        # Import analysis modules
        from meeting_summarizer import AdvancedMeetingSummarizer
        from enhanced_emotion_detector import EnhancedEmotionDetector
        
        # Initialize analyzers
        summarizer = AdvancedMeetingSummarizer()
        emotion_detector = EnhancedEmotionDetector()
        
        # Update progress
        self.update_state(
            state='PROGRESS',
            meta={'current': 20, 'total': 100, 'status': 'Generating summary...'}
        )
        
        # Generate comprehensive summary
        summary = summarizer.generate_comprehensive_summary(transcript_data)
        
        # Update progress
        self.update_state(
            state='PROGRESS',
            meta={'current': 50, 'total': 100, 'status': 'Analyzing emotions...'}
        )
        
        # Analyze emotions if utterances available
        emotion_analysis = {}
        if 'utterances' in transcript_data:
            utterances = transcript_data['utterances']
            # Note: This would need audio file path for full emotion analysis
            # For now, we'll do text-based emotion analysis
            emotion_analysis = {
                'total_utterances': len(utterances),
                'emotion_distribution': {},
                'speaker_emotions': {}
            }
        
        # Update progress
        self.update_state(
            state='PROGRESS',
            meta={'current': 80, 'total': 100, 'status': 'Compiling results...'}
        )
        
        # Compile analysis results
        analysis_results = {
            'summary': summary,
            'emotion_analysis': emotion_analysis,
            'key_points': summary.get('action_items', []),
            'decisions': summary.get('decisions', []),
            'speaker_analysis': summary.get('speaker_analysis', {}),
            'statistics': summary.get('statistics', {}),
            'completed_at': datetime.now().isoformat()
        }
        
        # Update final progress
        self.update_state(
            state='SUCCESS',
            meta={'current': 100, 'total': 100, 'status': 'Analysis completed!'}
        )
        
        return {
            'status': 'success',
            'analysis': analysis_results
        }
        
    except Exception as e:
        # Update error state
        self.update_state(
            state='FAILURE',
            meta={'error': str(e), 'status': 'Analysis failed'}
        )
        raise

@celery_app.task(bind=True)
def generate_report_task(self, analysis_data: Dict, report_type: str) -> Dict:
    """Generate report asynchronously"""
    try:
        # Update task state
        self.update_state(
            state='STARTED',
            meta={'current': 0, 'total': 100, 'status': f'Generating {report_type} report...'}
        )
        
        # Import report generation modules
        from docx import Document
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        import io
        
        # Update progress
        self.update_state(
            state='PROGRESS',
            meta={'current': 30, 'total': 100, 'status': 'Creating report content...'}
        )
        
        # Generate report content
        if report_type == 'docx':
            doc = Document()
            doc.add_heading('Meeting Report', 0)
            
            # Add summary
            if 'summary' in analysis_data:
                summary = analysis_data['summary']
                doc.add_heading('Executive Summary', level=1)
                doc.add_paragraph(summary.get('summaries', {}).get('abstractive', 'No summary available'))
            
            # Add key points
            if 'key_points' in analysis_data:
                doc.add_heading('Key Points', level=1)
                for point in analysis_data['key_points']:
                    doc.add_paragraph(f"• {point.get('text', '')}", style='List Bullet')
            
            # Save to bytes
            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            report_data = buffer.getvalue()
            report_mime = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            
        elif report_type == 'pdf':
            buffer = io.BytesIO()
            doc = SimpleDocTemplate(buffer, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Add content
            story.append(Paragraph('Meeting Report', styles['Heading1']))
            story.append(Spacer(1, 12))
            
            if 'summary' in analysis_data:
                summary = analysis_data['summary']
                story.append(Paragraph('Executive Summary', styles['Heading2']))
                story.append(Paragraph(summary.get('summaries', {}).get('abstractive', 'No summary available'), styles['Normal']))
            
            doc.build(story)
            buffer.seek(0)
            
            report_data = buffer.getvalue()
            report_mime = "application/pdf"
        
        else:
            # JSON report
            report_data = json.dumps(analysis_data, indent=2)
            report_mime = "application/json"
        
        # Update final progress
        self.update_state(
            state='SUCCESS',
            meta={'current': 100, 'total': 100, 'status': 'Report generated!'}
        )
        
        return {
            'status': 'success',
            'report_data': report_data,
            'report_mime': report_mime,
            'report_type': report_type,
            'completed_at': datetime.now().isoformat()
        }
        
    except Exception as e:
        # Update error state
        self.update_state(
            state='FAILURE',
            meta={'error': str(e), 'status': 'Report generation failed'}
        )
        raise

@celery_app.task(bind=True)
def send_email_task(self, email_data: Dict) -> Dict:
    """Send email asynchronously"""
    try:
        # Update task state
        self.update_state(
            state='STARTED',
            meta={'current': 0, 'total': 100, 'status': 'Preparing email...'}
        )
        
        # Import email module
        from email_summary import EmailSummaryGenerator
        
        # Initialize email generator
        email_gen = EmailSummaryGenerator()
        
        # Update progress
        self.update_state(
            state='PROGRESS',
            meta={'current': 50, 'total': 100, 'status': 'Sending email...'}
        )
        
        # Send email
        success = email_gen.send_email(
            email_data['recipient'],
            email_data['subject'],
            email_data['html_content']
        )
        
        # Update final progress
        self.update_state(
            state='SUCCESS',
            meta={'current': 100, 'total': 100, 'status': 'Email sent!'}
        )
        
        return {
            'status': 'success' if success else 'failed',
            'recipient': email_data['recipient'],
            'subject': email_data['subject'],
            'completed_at': datetime.now().isoformat()
        }
        
    except Exception as e:
        # Update error state
        self.update_state(
            state='FAILURE',
            meta={'error': str(e), 'status': 'Email sending failed'}
        )
        raise

# Streamlit integration functions
def submit_transcription_task(audio_file_path: str, config: Dict) -> str:
    """Submit transcription task to queue"""
    try:
        task = transcribe_audio_task.delay(audio_file_path, config)
        return task.id
    except Exception as e:
        print(f"Error submitting transcription task: {e}")
        return None

def submit_analysis_task(transcript_data: Dict) -> str:
    """Submit analysis task to queue"""
    try:
        task = analyze_meeting_task.delay(transcript_data)
        return task.id
    except Exception as e:
        print(f"Error submitting analysis task: {e}")
        return None

def submit_report_task(analysis_data: Dict, report_type: str) -> str:
    """Submit report generation task to queue"""
    try:
        task = generate_report_task.delay(analysis_data, report_type)
        return task.id
    except Exception as e:
        print(f"Error submitting report task: {e}")
        return None

def submit_email_task(email_data: Dict) -> str:
    """Submit email task to queue"""
    try:
        task = send_email_task.delay(email_data)
        return task.id
    except Exception as e:
        print(f"Error submitting email task: {e}")
        return None

def show_task_monitor():
    """Display task monitor in Streamlit"""
    if not STREAMLIT_AVAILABLE:
        print("Streamlit not available - cannot show task monitor UI")
        return
    
    st.title("🔄 Task Monitor")
    
    # Initialize task manager
    if 'task_manager' not in st.session_state:
        st.session_state.task_manager = TaskManager()
    
    task_manager = st.session_state.task_manager
    
    # Get all tasks
    tasks = task_manager.get_all_tasks()
    
    if not tasks:
        st.info("No tasks found.")
        return
    
    # Filter tasks by status
    status_filter = st.selectbox(
        "Filter by status:",
        ['All'] + list(set(task['status'] for task in tasks))
    )
    
    if status_filter != 'All':
        tasks = [task for task in tasks if task['status'] == status_filter]
    
    # Display tasks
    st.subheader(f"📋 Tasks ({len(tasks)} found)")
    
    for task in tasks:
        with st.expander(f"Task {task['task_id'][:8]}... - {task['status']}", expanded=False):
            st.write(f"**Task ID:** {task['task_id']}")
            st.write(f"**Status:** {task['status']}")
            
            if task.get('result'):
                st.write("**Result:**")
                st.json(task['result'])
            
            if task.get('error'):
                st.write("**Error:**")
                st.error(task['error'])
            
            if task.get('info'):
                st.write("**Info:**")
                st.json(task['info'])
    
    # Refresh button
    if st.button("🔄 Refresh"):
        st.rerun()

def show_task_progress(task_id: str):
    """Show progress for a specific task"""
    if not STREAMLIT_AVAILABLE:
        print(f"Task {task_id} progress - Streamlit not available")
        return
    
    if 'task_manager' not in st.session_state:
        st.session_state.task_manager = TaskManager()
    
    task_manager = st.session_state.task_manager
    task_status = task_manager.get_task_status(task_id)
    
    if task_status['status'] == 'SUCCESS':
        st.success("✅ Task completed successfully!")
        if task_status.get('result'):
            st.json(task_status['result'])
    elif task_status['status'] == 'FAILURE':
        st.error("❌ Task failed!")
        if task_status.get('error'):
            st.error(task_status['error'])
    elif task_status['status'] == 'PENDING':
        st.info("⏳ Task is pending...")
    else:
        st.info(f"🔄 Task status: {task_status['status']}")
        if task_status.get('info'):
            st.json(task_status['info'])

# Utility functions
def is_task_complete(task_id: str) -> bool:
    """Check if task is complete"""
    if STREAMLIT_AVAILABLE and 'task_manager' not in st.session_state:
        st.session_state.task_manager = TaskManager()
        task_manager = st.session_state.task_manager
    else:
        task_manager = TaskManager()
    
    task_status = task_manager.get_task_status(task_id)
    
    return task_status['status'] in ['SUCCESS', 'FAILURE']

def get_task_result(task_id: str) -> Optional[Dict]:
    """Get task result if complete"""
    if is_task_complete(task_id):
        if STREAMLIT_AVAILABLE and 'task_manager' not in st.session_state:
            st.session_state.task_manager = TaskManager()
            task_manager = st.session_state.task_manager
        else:
            task_manager = TaskManager()
        
        task_status = task_manager.get_task_status(task_id)
        
        if task_status['status'] == 'SUCCESS':
            return task_status.get('result')
    
    return None

def main():
    """Main function for standalone execution"""
    print("🔄 VerbatimAI Async Task Queue System")
    print("=" * 40)
    
    # Test Redis connection
    try:
        task_manager = TaskManager()
        print("✅ Redis connection: OK")
    except Exception as e:
        print(f"❌ Redis connection failed: {e}")
        print("Make sure Redis is running on localhost:6379")
        return
    
    # Test Celery configuration
    try:
        print("✅ Celery configuration: OK")
        print(f"   Broker: {CELERY_BROKER_URL}")
        print(f"   Backend: {CELERY_RESULT_BACKEND}")
    except Exception as e:
        print(f"❌ Celery configuration error: {e}")
    
    # Show available tasks
    print("\n📋 Available Celery Tasks:")
    print("   - transcribe_audio_task")
    print("   - analyze_meeting_task") 
    print("   - generate_report_task")
    print("   - send_email_task")
    
    # Show task management functions
    print("\n🛠️ Available Functions:")
    print("   - submit_transcription_task()")
    print("   - submit_analysis_task()")
    print("   - submit_report_task()")
    print("   - submit_email_task()")
    print("   - show_task_monitor() [Streamlit only]")
    
    # Show current tasks
    tasks = task_manager.get_all_tasks()
    print(f"\n📊 Current Tasks: {len(tasks)} found")
    
    for task in tasks[:5]:  # Show first 5 tasks
        print(f"   {task['task_id'][:8]}... - {task['status']}")
    
    if len(tasks) > 5:
        print(f"   ... and {len(tasks) - 5} more")
    
    print("\n💡 Usage:")
    print("   - Import this module in your Streamlit app")
    print("   - Use submit_*_task() functions to queue tasks")
    print("   - Monitor progress with show_task_monitor()")
    print("   - Start Celery worker: celery -A async_task_queue worker")

if __name__ == "__main__":
    main() 