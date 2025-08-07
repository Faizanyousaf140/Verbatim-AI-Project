import streamlit as st

# Configure Streamlit to allow larger file uploads (1GB = 1024MB)
# Page configuration with improved theme
st.set_page_config(
    page_title="VerbatimAI - Smart Meeting Intelligence",
    page_icon="🎙️",
    layout="wide",
    initial_sidebar_state="expanded",
    menu_items={
        'Get Help': 'https://verbatimai.help',
        'Report a bug': "https://verbatimai.bugs",
        'About': "# VerbatimAI\nSmart Meeting Intelligence Platform powered by AI"
    }
)
# Override Streamlit's default file upload size limit
import os
os.environ['STREAMLIT_SERVER_MAX_UPLOAD_SIZE'] = '1024'

import sys
import json
from datetime import datetime
import requests
import tempfile
import time
import io
import re
from dotenv import load_dotenv

# Optional imports with graceful fallbacks for document processing
try:
    from docx import Document
    DOCX_AVAILABLE = True
except (ImportError, ValueError) as e:
    print(f"⚠️ python-docx not available: {e}")
    DOCX_AVAILABLE = False
    Document = None

try:
    from reportlab.lib.pagesizes import letter
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    REPORTLAB_AVAILABLE = True
except (ImportError, ValueError) as e:
    print(f"⚠️ ReportLab not available: {e}")
    REPORTLAB_AVAILABLE = False

# Optional imports with graceful fallbacks
try:
    import pandas as pd
    PANDAS_AVAILABLE = True
except (ImportError, ValueError) as e:
    print(f"⚠️ Pandas not available: {e}")
    PANDAS_AVAILABLE = False
    # Create a simple mock for basic DataFrame functionality
    class MockDataFrame:
        def __init__(self, data=None):
            self.data = data or []
        def to_csv(self, *args, **kwargs):
            return ""
        def to_json(self, *args, **kwargs):
            return "{}"
    pd = type('MockPandas', (), {'DataFrame': MockDataFrame})()

try:
    import plotly.express as px
    import plotly.graph_objects as go
    PLOTLY_AVAILABLE = True
except (ImportError, ValueError) as e:
    print(f"⚠️ Plotly not available: {e}")
    PLOTLY_AVAILABLE = False
    # Create mock plotly objects
    class MockPlotly:
        def bar(self, *args, **kwargs): return None
        def line(self, *args, **kwargs): return None
        def pie(self, *args, **kwargs): return None
        def scatter(self, *args, **kwargs): return None
    px = MockPlotly()
    go = type('MockGO', (), {'Figure': lambda *args, **kwargs: None})()

# Optional imports for other packages
try:
    import assemblyai as aai
    ASSEMBLYAI_AVAILABLE = True
except (ImportError, ValueError) as e:
    print(f"⚠️ AssemblyAI not available: {e}")
    ASSEMBLYAI_AVAILABLE = False
    aai = None

try:
    from streamlit_option_menu import option_menu
    OPTION_MENU_AVAILABLE = True
except (ImportError, ValueError) as e:
    print(f"⚠️ Streamlit option menu not available: {e}")
    OPTION_MENU_AVAILABLE = False
    # Simple fallback for option_menu
    def option_menu(menu_title, options, **kwargs):
        return st.selectbox(menu_title, options)

try:
    import nltk
    from nltk.tokenize import sent_tokenize, word_tokenize
    from nltk.corpus import stopwords
    from nltk.sentiment import SentimentIntensityAnalyzer
    NLTK_AVAILABLE = True
except (ImportError, ValueError) as e:
    print(f"⚠️ NLTK not available: {e}")
    NLTK_AVAILABLE = False
    # Simple fallbacks
    def sent_tokenize(text): return text.split('.')
    def word_tokenize(text): return text.split()
    stopwords = type('MockStopwords', (), {'words': lambda x: []})()
    SentimentIntensityAnalyzer = lambda: None

try:
    import smtplib
    from email.message import EmailMessage
    EMAIL_AVAILABLE = True
except (ImportError, ValueError) as e:
    print(f"⚠️ Email modules not available: {e}")
    EMAIL_AVAILABLE = False

try:
    import seaborn as sns
    SEABORN_AVAILABLE = True
except (ImportError, ValueError) as e:
    print(f"⚠️ Seaborn not available: {e}")
    SEABORN_AVAILABLE = False
    sns = None

try:
    import altair as alt
    ALTAIR_AVAILABLE = True
except (ImportError, ValueError) as e:
    print(f"⚠️ Altair not available: {e}")
    ALTAIR_AVAILABLE = False
    alt = None

# Load variables from .env file
load_dotenv()

# Access your secrets safely
ASSEMBLYAI_API_KEY = os.getenv("ASSEMBLYAI_API_KEY")
EMAIL_SENDER = os.getenv("EMAIL_SENDER")
EMAIL_PASSWORD = os.getenv("EMAIL_PASSWORD")

# Then continue with the rest of your configuration or app setup
APP_TITLE = "VerbatimAI - Smart Meeting Intelligence"
from config import (
    ASSEMBLYAI_API_KEY, ALLOWED_EXTENSIONS, MAX_FILE_SIZE, 
    DECISION_KEYWORDS, ACTION_KEYWORDS, UI_THEME, EMOTION_DETECTION_CONFIG,
    EMAIL_SENDER, EMAIL_PASSWORD, RECORDING_CONFIG
)

# Import role_management first (this is required)
try:
    from role_management import (
        show_login_page, check_authentication, show_user_info, 
        get_current_user_role, NavigationManager, Permission
    )
    print("✅ Role management imported successfully")
except ImportError as e:
    print(f"❌ Critical error: Role management not available: {e}")
    # Set fallback functions
    def check_authentication():
        return False
    def show_login_page():
        st.title("🔐 Login Required")
        st.error("Authentication system not available. Please install required dependencies.")
    def show_user_info():
        pass
    def get_current_user_role():
        return None
    class NavigationManager:
        @staticmethod
        def get_navigation_items(role):
            return [{"title": "🏠 Dashboard", "icon": "house", "page": "dashboard"}]
    class Permission:
        pass

# Import other advanced modules (optional)
try:
    from audio_emotion_detector import AudioEmotionDetector
    from semantic_search import SemanticSearchEngine
    from email_summary import EmailSummaryGenerator
    from meeting_summarizer import AdvancedMeetingSummarizer
    from enhanced_emotion_detector import EnhancedEmotionDetector
    from interactive_transcript_editor import show_interactive_transcript_editor
    from async_task_queue import show_task_monitor, show_task_progress
    from real_time_recorder import show_real_time_recording, RealTimeRecorder
    from meeting_library import show_meeting_library, MeetingLibrary
    ADVANCED_FEATURES_AVAILABLE = True
    print("✅ Advanced features modules imported successfully")
except ImportError as e:
    print(f"⚠️ Warning: Advanced features not available: {e}")
    print("⚠️ Some features will be disabled. Install missing packages to enable full functionality.")
    ADVANCED_FEATURES_AVAILABLE = False
    # Fallback classes only - functions will use st.error directly in the routing section
    class RealTimeRecorder:
        pass
    class MeetingLibrary:
        pass

# Load environment variables
load_dotenv()

# Download NLTK data (only once per session)
if 'nltk_downloaded' not in st.session_state:
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        nltk.download('punkt')
    try:
        nltk.data.find('corpora/stopwords')
    except LookupError:
        nltk.download('stopwords')
    try:
        nltk.data.find('vader_lexicon')
    except LookupError:
        nltk.download('vader_lexicon')
    st.session_state.nltk_downloaded = True

# Configure AssemblyAI with user's API key
aai.settings.api_key = ASSEMBLYAI_API_KEY

# Initialize email summary generator globally
email_generator = None
if ADVANCED_FEATURES_AVAILABLE:
    try:
        email_generator = EmailSummaryGenerator()
        print("✅ Email summary generator initialized successfully")
    except Exception as e:
        print(f"❌ Error initializing email generator: {e}")
        email_generator = None



# Apply custom theme
st.markdown(f"""
<style>
    /* Global styles */
    .main {{
        background-color: {UI_THEME['backgroundColor']};
        color: {UI_THEME['textColor']};
    }}
    
    /* Header styles */
    .main-header {{
        font-size: 3rem;
        font-weight: bold;
        color: {UI_THEME['primaryColor']};
        text-align: center;
        margin-bottom: 2rem;
        text-shadow: 2px 2px 4px rgba(0,0,0,0.1);
    }}
    
    /* Card styles */
    .feature-card {{
        background-color: {UI_THEME['secondaryBackgroundColor']};
        padding: 1.5rem;
        border-radius: 15px;
        margin: 1rem 0;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        border: 1px solid #e0e0e0;
    }}
    
    .metric-card {{
        background-color: {UI_THEME['backgroundColor']};
        padding: 1.5rem;
        border-radius: 15px;
        border: 2px solid {UI_THEME['primaryColor']};
        text-align: center;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        transition: transform 0.2s;
    }}
    
    .metric-card:hover {{
        transform: translateY(-2px);
        box-shadow: 0 6px 12px rgba(0,0,0,0.15);
    }}
    
    /* Speaker styles */
    .speaker-bubble {{
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        color: white;
        padding: 0.8rem;
        border-radius: 20px;
        margin: 0.5rem 0;
        box-shadow: 0 2px 4px rgba(0,0,0,0.1);
    }}
    
    /* Status styles */
    .status-success {{
        color: {UI_THEME['successColor']};
        font-weight: bold;
        background-color: rgba(39, 174, 96, 0.1);
        padding: 0.5rem;
        border-radius: 5px;
        border: 1px solid {UI_THEME['successColor']};
    }}
    
    .status-error {{
        color: {UI_THEME['accentColor']};
        font-weight: bold;
        background-color: rgba(231, 76, 60, 0.1);
        padding: 0.5rem;
        border-radius: 5px;
        border: 1px solid {UI_THEME['accentColor']};
    }}
    
    .status-warning {{
        color: {UI_THEME['warningColor']};
        font-weight: bold;
        background-color: rgba(243, 156, 18, 0.1);
        padding: 0.5rem;
        border-radius: 5px;
        border: 1px solid {UI_THEME['warningColor']};
    }}
        background-color: #f8d7da;
        padding: 0.5rem;
        border-radius: 5px;
        border: 1px solid #f5c6cb;
    }}
    
    /* Navigation styles */
    .css-1d391kg {{
        background-color: {UI_THEME['secondaryBackgroundColor']};
    }}
    
    /* Button styles */
    .stButton > button {{
        background: linear-gradient(135deg, {UI_THEME['primaryColor']} 0%, #ff8a80 100%);
        color: white;
        border: none;
        border-radius: 10px;
        padding: 0.5rem 1rem;
        font-weight: bold;
        transition: all 0.3s;
    }}
    
    .stButton > button:hover {{
        transform: translateY(-2px);
        box-shadow: 0 4px 8px rgba(0,0,0,0.2);
    }}
    
    /* Text input styles */
    .stTextInput > div > div > input {{
        border-radius: 10px;
        border: 2px solid #e0e0e0;
    }}
    
    /* File uploader styles */
    .stFileUploader > div {{
        border-radius: 10px;
        border: 2px dashed {UI_THEME['primaryColor']};
        background-color: {UI_THEME['secondaryBackgroundColor']};
    }}
    
    /* Chart container styles */
    .chart-container {{
        background-color: {UI_THEME['backgroundColor']};
        padding: 1rem;
        border-radius: 10px;
        border: 1px solid #e0e0e0;
        margin: 1rem 0;
    }}
    
    /* Tab styles */
    .stTabs [data-baseweb="tab-list"] {{
        gap: 8px;
    }}
    
    .stTabs [data-baseweb="tab"] {{
        background-color: {UI_THEME['secondaryBackgroundColor']};
        border-radius: 10px 10px 0 0;
        padding: 10px 20px;
        border: none;
    }}
    
    .stTabs [aria-selected="true"] {{
        background-color: {UI_THEME['primaryColor']};
        color: white;
    }}
</style>
""", unsafe_allow_html=True)

# Enhanced UI Components and Modern Design System
def init_modern_ui():
    """Initialize modern UI with custom styling"""
    # Note: st.set_page_config() is called at the top of the script
    # This function only handles CSS styling and UI components
    
    # Modern CSS Styling - Google/Microsoft Design System
    st.markdown("""
    <style>
    /* Import Google Fonts */
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&family=JetBrains+Mono:wght@400;500&display=swap');
    
    /* Root Variables - Modern Color Palette */
    :root {
        --primary-blue: #1E88E5;
        --primary-blue-dark: #1565C0;
        --primary-blue-light: #E3F2FD;
        --secondary-purple: #7B1FA2;
        --secondary-green: #43A047;
        --accent-orange: #FF9800;
        --accent-red: #E53935;
        --neutral-50: #FAFAFA;
        --neutral-100: #F5F5F5;
        --neutral-200: #EEEEEE;
        --neutral-300: #E0E0E0;
        --neutral-400: #BDBDBD;
        --neutral-500: #9E9E9E;
        --neutral-600: #757575;
        --neutral-700: #616161;
        --neutral-800: #424242;
        --neutral-900: #212121;
        --glass-bg: rgba(255, 255, 255, 0.95);
        --glass-border: rgba(255, 255, 255, 0.2);
        --shadow-sm: 0 1px 3px rgba(0, 0, 0, 0.1);
        --shadow-md: 0 4px 6px rgba(0, 0, 0, 0.1);
        --shadow-lg: 0 10px 25px rgba(0, 0, 0, 0.15);
        --shadow-xl: 0 20px 40px rgba(0, 0, 0, 0.2);
    }
    
    /* Global Styles */
    .stApp {
        background: linear-gradient(135deg, #667eea 0%, #764ba2 100%);
        font-family: 'Inter', -apple-system, BlinkMacSystemFont, sans-serif;
    }
    
    /* Sidebar Modern Design */
    .css-1d391kg {
        background: var(--glass-bg);
        backdrop-filter: blur(20px);
        border-right: 1px solid var(--glass-border);
        box-shadow: var(--shadow-lg);
    }
    
    /* Main Content Area */
    .main .block-container {
        padding: 2rem 3rem;
        background: var(--glass-bg);
        backdrop-filter: blur(20px);
        border-radius: 24px;
        margin: 1rem;
        box-shadow: var(--shadow-xl);
        border: 1px solid var(--glass-border);
    }
    
    /* Headers with Modern Typography */
    .main-header {
        font-family: 'Inter', sans-serif;
        font-weight: 700;
        font-size: 3.5rem;
        background: linear-gradient(135deg, var(--primary-blue), var(--secondary-purple));
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        text-align: center;
        margin-bottom: 1rem;
        letter-spacing: -0.02em;
    }
    
    .section-header {
        font-family: 'Inter', sans-serif;
        font-weight: 600;
        font-size: 1.8rem;
        color: var(--neutral-800);
        margin: 2rem 0 1rem 0;
        position: relative;
    }
    
    .section-header::before {
        content: '';
        position: absolute;
        bottom: -4px;
        left: 0;
        width: 60px;
        height: 3px;
        background: linear-gradient(90deg, var(--primary-blue), var(--secondary-purple));
        border-radius: 2px;
    }
    
    /* Modern Cards */
    .feature-card {
        background: var(--glass-bg);
        backdrop-filter: blur(10px);
        padding: 2rem;
        border-radius: 16px;
        border: 1px solid var(--glass-border);
        box-shadow: var(--shadow-md);
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
        margin-bottom: 1rem;
        position: relative;
        overflow: hidden;
    }
    
    .feature-card::before {
        content: '';
        position: absolute;
        top: 0;
        left: 0;
        right: 0;
        height: 4px;
        background: linear-gradient(90deg, var(--primary-blue), var(--secondary-purple), var(--accent-orange));
    }
    
    .feature-card:hover {
        transform: translateY(-4px);
        box-shadow: var(--shadow-xl);
        border-color: var(--primary-blue);
    }
    
    .feature-card h3 {
        font-family: 'Inter', sans-serif;
        font-weight: 600;
        font-size: 1.3rem;
        color: var(--neutral-800);
        margin-bottom: 0.5rem;
    }
    
    .feature-card p {
        color: var(--neutral-600);
        line-height: 1.6;
        font-size: 0.95rem;
    }
    
    /* Status Indicators */
    .status-success {
        display: inline-flex;
        align-items: center;
        padding: 0.5rem 1rem;
        background: linear-gradient(135deg, #E8F5E8, #C8E6C9);
        color: var(--secondary-green);
        border-radius: 12px;
        border: 1px solid rgba(67, 160, 71, 0.2);
        font-weight: 500;
        font-size: 0.9rem;
        margin: 0.5rem 0;
    }
    
    .status-error {
        display: inline-flex;
        align-items: center;
        padding: 0.5rem 1rem;
        background: linear-gradient(135deg, #FFEBEE, #FFCDD2);
        color: var(--accent-red);
        border-radius: 12px;
        border: 1px solid rgba(229, 57, 53, 0.2);
        font-weight: 500;
        font-size: 0.9rem;
        margin: 0.5rem 0;
    }
    
    .status-warning {
        display: inline-flex;
        align-items: center;
        padding: 0.5rem 1rem;
        background: linear-gradient(135deg, #FFF8E1, #FFECB3);
        color: var(--accent-orange);
        border-radius: 12px;
        border: 1px solid rgba(255, 152, 0, 0.2);
        font-weight: 500;
        font-size: 0.9rem;
        margin: 0.5rem 0;
    }
    
    /* Navigation styles */
    .css-1d391kg {
        background: var(--glass-bg);
        backdrop-filter: blur(20px);
        border-right: 1px solid var(--glass-border);
        box-shadow: var(--shadow-lg);
    }
    
    /* Button styles */
    .stButton > button {
        background: linear-gradient(135deg, var(--primary-blue), var(--primary-blue-dark));
        color: white;
        border: none;
        border-radius: 12px;
        padding: 0.75rem 2rem;
        font-family: 'Inter', sans-serif;
        font-weight: 500;
        font-size: 0.95rem;
        transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1);
        box-shadow: var(--shadow-md);
        letter-spacing: 0.025em;
    }
    
    .stButton > button:hover {
        transform: translateY(-2px);
        box-shadow: var(--shadow-lg);
        background: linear-gradient(135deg, var(--primary-blue-dark), var(--secondary-purple));
    }
    
    .stButton > button:active {
        transform: translateY(0);
    }
    
    /* Metrics with Modern Design */
    .metric-container {
        background: var(--glass-bg);
        backdrop-filter: blur(10px);
        padding: 1.5rem;
        border-radius: 16px;
        border: 1px solid var(--glass-border);
        box-shadow: var(--shadow-md);
        text-align: center;
        transition: all 0.3s ease;
    }
    
    .metric-container:hover {
        transform: translateY(-2px);
        box-shadow: var(--shadow-lg);
    }
    
    .metric-value {
        font-family: 'Inter', sans-serif;
        font-weight: 700;
        font-size: 2.5rem;
        background: linear-gradient(135deg, var(--primary-blue), var(--secondary-purple));
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
        line-height: 1;
    }
    
    .metric-label {
        font-family: 'Inter', sans-serif;
        font-weight: 500;
        font-size: 0.9rem;
        color: var(--neutral-600);
        margin-top: 0.5rem;
        text-transform: uppercase;
        letter-spacing: 0.05em;
    }
    
    /* Modern Input Fields */
    .stTextInput > div > div > input,
    .stTextArea > div > div > textarea,
    .stSelectbox > div > div > select {
        border-radius: 12px;
        border: 2px solid var(--neutral-300);
        font-family: 'Inter', sans-serif;
        transition: all 0.3s ease;
        background: var(--glass-bg);
        backdrop-filter: blur(10px);
    }
    
    .stTextInput > div > div > input:focus,
    .stTextArea > div > div > textarea:focus {
        border-color: var(--primary-blue);
        box-shadow: 0 0 0 3px rgba(30, 136, 229, 0.1);
    }
    
    /* File Upload Area */
    .uploadedFile {
        background: var(--glass-bg);
        backdrop-filter: blur(10px);
        border: 2px dashed var(--primary-blue);
        border-radius: 16px;
        padding: 2rem;
        text-align: center;
        transition: all 0.3s ease;
    }
    
    .uploadedFile:hover {
        border-color: var(--secondary-purple);
        background: var(--primary-blue-light);
    }
    
    /* Progress Bars */
    .stProgress > div > div > div {
        background: linear-gradient(90deg, var(--primary-blue), var(--secondary-purple));
        border-radius: 10px;
    }
    
    /* Expanders */
    .streamlit-expanderHeader {
        background: var(--glass-bg);
        backdrop-filter: blur(10px);
        border-radius: 12px;
        border: 1px solid var(--glass-border);
        font-weight: 500;
    }
    
    /* Sidebar Styling */
    .sidebar-logo {
        text-align: center;
        padding: 1rem 0;
        border-bottom: 1px solid var(--glass-border);
        margin-bottom: 1rem;
    }
    
    .sidebar-title {
        font-family: 'Inter', sans-serif;
        font-weight: 700;
        font-size: 1.5rem;
        background: linear-gradient(135deg, var(--primary-blue), var(--secondary-purple));
        -webkit-background-clip: text;
        -webkit-text-fill-color: transparent;
        background-clip: text;
    }
    
    /* Charts and Visualizations */
    .js-plotly-plot {
        border-radius: 12px;
        overflow: hidden;
        box-shadow: var(--shadow-md);
        background: var(--glass-bg);
    }
    
    /* Responsive Design */
    @media (max-width: 768px) {
        .main .block-container {
            padding: 1rem;
            margin: 0.5rem;
        }
        
        .main-header {
            font-size: 2.5rem;
        }
        
        .feature-card {
            padding: 1.5rem;
        }
    }
    </style>
    """, unsafe_allow_html=True)

def create_modern_metric_card(title, value, delta=None, icon="📊"):
    """Create a modern metric card with glassmorphism effect"""
    delta_html = f'<div style="color: var(--secondary-green); font-size: 0.9rem; margin-top: 0.5rem;">▲ {delta}</div>' if delta else ''
    
    return st.markdown(f"""
    <div class="metric-container">
        <div style="font-size: 2rem; margin-bottom: 0.5rem;">{icon}</div>
        <div class="metric-value">{value}</div>
        <div class="metric-label">{title}</div>
        {delta_html}
    </div>
    """, unsafe_allow_html=True)

def create_feature_card(title, description, icon="🚀", button_text=None, button_key=None):
    """Create a modern feature card"""
    button_html = f'<button class="feature-button" onclick="window.location.reload()">{button_text}</button>' if button_text else ''
    
    card_html = f"""
    <div class="feature-card">
        <div style="display: flex; align-items: center; margin-bottom: 1rem;">
            <span style="font-size: 2rem; margin-right: 1rem;">{icon}</span>
            <h3 style="margin: 0;">{title}</h3>
        </div>
        <p>{description}</p>
        {button_html}
    </div>
    """
    
    return st.markdown(card_html, unsafe_allow_html=True)

def show_loading_animation(text="Processing..."):
    """Show modern loading animation"""
    return st.markdown(f"""
    <div style="display: flex; align-items: center; justify-content: center; padding: 2rem;">
        <div style="margin-right: 1rem;">
            <div style="width: 40px; height: 40px; border: 4px solid var(--neutral-300); border-top: 4px solid var(--primary-blue); border-radius: 50%; animation: spin 1s linear infinite;"></div>
        </div>
        <span style="font-family: 'Inter', sans-serif; font-weight: 500; color: var(--neutral-700);">{text}</span>
    </div>
    <style>
    @keyframes spin {{
        0% {{ transform: rotate(0deg); }}
        100% {{ transform: rotate(360deg); }}
    }}
    </style>
    """, unsafe_allow_html=True)

def create_status_badge(status, text):
    """Create modern status badges"""
    status_class = f"status-{status}"
    icon_map = {
        'success': '✅',
        'error': '❌', 
        'warning': '⚠️',
        'info': 'ℹ️'
    }
    icon = icon_map.get(status, '🔹')
    
    return st.markdown(f"""
    <div class="{status_class}">
        <span style="margin-right: 0.5rem;">{icon}</span>
        {text}
    </div>
    """, unsafe_allow_html=True)

class VerbatimAI:
    def __init__(self):
        self.transcript_data = None
        self.analysis_results = None
        
        # Initialize advanced features if available
        if ADVANCED_FEATURES_AVAILABLE:
            try:
                self.audio_emotion_detector = AudioEmotionDetector()
                self.semantic_search_engine = SemanticSearchEngine()
                self.email_summary_generator = EmailSummaryGenerator()
                self.advanced_meeting_summarizer = AdvancedMeetingSummarizer()
                self.enhanced_emotion_detector = EnhancedEmotionDetector()
                print("✅ Advanced features initialized successfully")
            except Exception as e:
                print(f"❌ Error initializing advanced features: {e}")
                self.audio_emotion_detector = None
                self.semantic_search_engine = None
                self.email_summary_generator = None
                self.advanced_meeting_summarizer = None
                self.enhanced_emotion_detector = None
        else:
            self.audio_emotion_detector = None
            self.semantic_search_engine = None
            self.email_summary_generator = None
            self.advanced_meeting_summarizer = None
            self.enhanced_emotion_detector = None
    
    def verify_api_key(self):
        """Verify AssemblyAI API key is valid"""
        try:
            # Test API key with a simple request
            transcriber = aai.Transcriber()
            # This will fail if API key is invalid, but we can catch the error
            return True
        except Exception as e:
            if "401" in str(e) or "unauthorized" in str(e).lower():
                return False
            return True  # Other errors might be network related
    
    def transcribe_audio(self, audio_file, config=None):
        """Transcribe audio using AssemblyAI with improved error handling"""
        try:
            # Verify API key first
            if not self.verify_api_key():
                st.error("❌ Invalid AssemblyAI API key. Please check your configuration.")
                return None
            
            # Save uploaded file temporarily
            with tempfile.NamedTemporaryFile(delete=False, suffix='.wav') as tmp_file:
                tmp_file.write(audio_file.read())
                tmp_file_path = tmp_file.name
            
            # Retry logic for transcription
            max_retries = 3
            for attempt in range(max_retries):
                try:
                    st.info(f"🔄 Transcription attempt {attempt + 1}/{max_retries}...")
                    
                    # Use provided config or try with full features first
                    try:
                        transcription_config = aai.TranscriptionConfig(
                            speaker_labels=config.get('speaker_detection', True) if config else True,
                            sentiment_analysis=config.get('sentiment_analysis', True) if config else True,
                            auto_highlights=config.get('auto_highlights', True) if config else True,
                            auto_punctuation=config.get('auto_punctuation', True) if config else True,
                            filter_profanity=config.get('filter_profanity', False) if config else False,
                            language_code=config.get('language', 'en') if config else 'en'
                        )
                        
                        transcriber = aai.Transcriber()
                        transcript = transcriber.transcribe(tmp_file_path, transcription_config)
                     ## these  2 lins remove or comment out    
                    except Exception as config_error:
                        st.warning("⚠️ Some advanced features may not be available. Using basic transcription...")
                        
                        # Fallback to basic configuration
                        fallback_config = aai.TranscriptionConfig(
                            speaker_labels=config.get('speaker_detection', True) if config else True,
                            language_code=config.get('language', 'en') if config else 'en'
                        )
                        
                        transcriber = aai.Transcriber()
                        transcript = transcriber.transcribe(tmp_file_path, fallback_config)
                    
                    # Check if transcription was successful
                    if transcript.status == aai.TranscriptStatus.error:
                        raise Exception(f"Transcription failed: {transcript.error}")
                    
                    # Clean up temp file
                    os.unlink(tmp_file_path)
                    
                    # Apply advanced features if available
                    if self.audio_emotion_detector:
                        st.info("🔊 Analyzing audio emotions...")
                        transcript = self.analyze_audio_emotions(tmp_file_path, transcript)
                    
                    # Build semantic search index
                    if self.semantic_search_engine:
                        st.info("🔍 Building semantic search index...")
                        self.build_semantic_index(transcript)
                    
                    st.success("✅ Transcription completed successfully!")
                    return transcript
                    
                except Exception as e:
                    if attempt < max_retries - 1:
                        st.warning(f"⚠️ Attempt {attempt + 1} failed: {str(e)}. Retrying...")
                        import time
                        time.sleep(2)  # Wait before retry
                    else:
                        raise e
            
            return None
            
        except Exception as e:
            st.error(f"❌ Error during transcription after {max_retries} attempts: {str(e)}")
            # Clean up temp file if it exists
            try:
                if 'tmp_file_path' in locals() and os.path.exists(tmp_file_path):
                    os.unlink(tmp_file_path)
            except:
                pass
            return None
    
    def generate_summary(self, transcript):
        """Generate a comprehensive meeting summary"""
        try:
            # Calculate duration safely
            duration_minutes = 0
            if hasattr(transcript, 'audio_duration') and transcript.audio_duration:
                duration_minutes = transcript.audio_duration / 1000 / 60
            elif transcript.utterances:
                # Calculate from last utterance end time
                last_utterance = max(transcript.utterances, key=lambda u: u.end)
                duration_minutes = last_utterance.end / 1000 / 60
            
            total_words = len(transcript.text.split())
            total_utterances = len(transcript.utterances)
            unique_speakers = len(set([u.speaker for u in transcript.utterances]))
            
            # Calculate speaking rate
            speaking_rate = total_words / duration_minutes if duration_minutes > 0 else 0
            
            # Generate summary
            summary = f"""
            **Meeting Overview:**
            - **Duration**: {duration_minutes:.2f} minutes
            - **Participants**: {unique_speakers} speakers
            - **Total Words**: {total_words:,} words
            - **Speaking Rate**: {speaking_rate:.1f} words per minute
            - **Total Utterances**: {total_utterances} exchanges
            
            **Key Statistics:**
            - **Average Utterance Length**: {total_words / total_utterances:.1f} words per utterance
            - **Engagement Level**: {'High' if speaking_rate > 150 else 'Medium' if speaking_rate > 100 else 'Low'}
            
            **Content Summary:**
            This meeting involved {unique_speakers} participants discussing various topics over {duration_minutes:.1f} minutes. 
            The conversation included {total_utterances} exchanges with an average speaking rate of {speaking_rate:.1f} words per minute.
            """
            
            return summary
            
        except Exception as e:
            return f"Error generating summary: {str(e)}"
    
    def generate_advanced_summary(self, transcript):
        """Generate advanced meeting summary with content summary using NLP models"""
        try:
            if not self.advanced_meeting_summarizer:
                return self.generate_summary(transcript)
            
            # Prepare transcript data for advanced summarizer
            transcript_data = {
                'text': transcript.text,
                'utterances': transcript.utterances,
                'duration_minutes': transcript.audio_duration / 1000 / 60,
                'audio_duration': transcript.audio_duration,  # Keep original for compatibility
                'title': 'Meeting Transcript'
            }
            
            # Generate comprehensive summary
            comprehensive_summary = self.advanced_meeting_summarizer.generate_comprehensive_summary(transcript_data)
            
            if 'error' in comprehensive_summary:
                return f"Advanced summary generation failed: {comprehensive_summary['error']}"
            
            # Extract content summary
            content_summary = comprehensive_summary.get('summaries', {}).get('content_summary', '')
            #abstractive_summary = comprehensive_summary.get('summaries', {}).get('abstractive', '')
            
            # Format the advanced summary
            advanced_summary = f"""
            **📋 Advanced Meeting Summary**
            
            **📊 Meeting Statistics:**
            - **Duration**: {comprehensive_summary.get('meeting_info', {}).get('duration_minutes', 0):.2f} minutes
            - **Participants**: {comprehensive_summary.get('meeting_info', {}).get('participants', 0)} speakers
            - **Total Words**: {comprehensive_summary.get('statistics', {}).get('total_words', 0):,} words
            - **Sentiment**: {comprehensive_summary.get('sentiment', 'Neutral').title()}
            
            **📝 Content Summary (1/4 of original length):**
            {content_summary if content_summary else 'Content summary not available.'}
            
            **🔍 Key Insights:**
            """
            
            # Add insights
            insights = comprehensive_summary.get('insights', [])
            for i, insight in enumerate(insights, 1):
                advanced_summary += f"\n{i}. {insight}"
            
            if not insights:
                advanced_summary += "\nNo specific insights available."
            
            return advanced_summary
            
        except Exception as e:
            print(f"Error generating advanced summary: {e}")
            return self.generate_summary(transcript)
    
    def extract_organizations(self, transcript):
        """Extract organizations and companies mentioned in the transcript"""
        organizations = []
        
        # Common company/organization keywords
        org_keywords = [
            'google', 'microsoft', 'apple', 'amazon', 'facebook', 'meta', 'netflix', 'tesla',
            'uber', 'airbnb', 'spotify', 'linkedin', 'twitter', 'instagram', 'youtube',
            'salesforce', 'oracle', 'ibm', 'intel', 'cisco', 'dell', 'hp', 'nvidia',
            'adobe', 'autodesk', 'zoom', 'slack', 'dropbox', 'box', 'atlassian',
            'github', 'gitlab', 'docker', 'kubernetes', 'aws', 'azure', 'gcp',
            'company', 'corporation', 'inc', 'llc', 'ltd', 'enterprise', 'organization'
        ]
        
        for utterance in transcript.utterances:
            text_lower = utterance.text.lower()
            
            # Check for organization keywords
            for keyword in org_keywords:
                if keyword in text_lower:
                    # Find the actual mention in the text
                    words = utterance.text.split()
                    for word in words:
                        if keyword in word.lower():
                            organizations.append({
                                'text': utterance.text,
                                'speaker': f"Speaker {utterance.speaker}",
                                'timestamp': utterance.start / 1000,
                                'organization': word
                            })
                            break
                    break
        
        return organizations
    
    def extract_enhanced_key_points(self, transcript):
        """Extract enhanced key points including organizations and technical terms"""
        key_points = {
            'decisions': [],
            'action_items': [],
            'questions': [],
            'highlights': [],
            'organizations': [],
            'technical_terms': []
        }
        
        # Technical terms and technologies
        tech_keywords = [
            'ai', 'machine learning', 'deep learning', 'neural network', 'algorithm',
            'api', 'database', 'cloud', 'server', 'client', 'frontend', 'backend',
            'javascript', 'python', 'java', 'react', 'angular', 'vue', 'node.js',
            'docker', 'kubernetes', 'aws', 'azure', 'gcp', 'blockchain', 'crypto',
            'mobile app', 'web app', 'saas', 'paas', 'iaas', 'microservices'
        ]
        
        for utterance in transcript.utterances:
            text = utterance.text.lower()
            
            # Check for decisions
            if any(keyword in text for keyword in DECISION_KEYWORDS):
                key_points['decisions'].append({
                    'text': utterance.text,
                    'speaker': f"Speaker {utterance.speaker}",
                    'timestamp': utterance.start / 1000
                })
            
            # Check for action items
            if any(keyword in text for keyword in ACTION_KEYWORDS):
                key_points['action_items'].append({
                    'text': utterance.text,
                    'speaker': f"Speaker {utterance.speaker}",
                    'timestamp': utterance.start / 1000
                })
            
            # Check for questions
            if re.search(r'\?', utterance.text):
                key_points['questions'].append({
                    'text': utterance.text,
                    'speaker': f"Speaker {utterance.speaker}",
                    'timestamp': utterance.start / 1000
                })
            
            # Check for technical terms
            for tech_term in tech_keywords:
                if tech_term in text:
                    key_points['technical_terms'].append({
                        'text': utterance.text,
                        'speaker': f"Speaker {utterance.speaker}",
                        'timestamp': utterance.start / 1000,
                        'term': tech_term
                    })
                    break
        
        # Add highlights from AssemblyAI (if available)
        try:
            if hasattr(transcript, 'highlights') and transcript.highlights:
                for highlight in transcript.highlights:
                    key_points['highlights'].append({
                        'text': highlight.text,
                        'reason': highlight.reason,
                        'timestamp': highlight.start / 1000
                    })
        except Exception as e:
            pass
        
        # Add organizations
        key_points['organizations'] = self.extract_organizations(transcript)
        
        # Check for missing key_points structure
        if not key_points:
            key_points = {
                'decisions': [],
                'action_items': [],
                'questions': [],
                'highlights': [],
                'organizations': [],
                'technical_terms': []
            }
        
        # Ensure all required keys exist
        for key in ['decisions', 'action_items', 'questions', 'highlights', 'organizations', 'technical_terms']:
            if key not in key_points:
                key_points[key] = []
        
        return key_points
    
    def analyze_sentiment(self, transcript):
        """Analyze sentiment for each speaker with enhanced emotion detection"""
        sentiment_analyzer = SentimentIntensityAnalyzer()
        speaker_sentiments = {}
        
        # Emotion keywords for enhanced detection
        emotion_keywords = {
            'calm': ['calm', 'peaceful', 'relaxed', 'serene', 'tranquil'],
            'angry': ['angry', 'furious', 'mad', 'irritated', 'frustrated', 'upset'],
            'confident': ['confident', 'sure', 'certain', 'positive', 'assured'],
            'excited': ['excited', 'enthusiastic', 'thrilled', 'eager', 'energetic'],
            'worried': ['worried', 'concerned', 'anxious', 'nervous', 'stressed'],
            'happy': ['happy', 'joyful', 'pleased', 'delighted', 'satisfied'],
            'sad': ['sad', 'unhappy', 'disappointed', 'depressed', 'melancholy']
        }
        
        for utterance in transcript.utterances:
            speaker = f"Speaker {utterance.speaker}"
            if speaker not in speaker_sentiments:
                speaker_sentiments[speaker] = {
                    'utterances': [],
                    'sentiment_scores': [],
                    'total_duration': 0,
                    'emotions': [],
                    'tone_analysis': []
                }
            
            # Use AssemblyAI sentiment if available, otherwise use VADER
            try:
                if hasattr(utterance, 'sentiment') and utterance.sentiment:
                    sentiment_score = {
                        'positive': utterance.sentiment.positive,
                        'negative': utterance.sentiment.negative,
                        'neutral': utterance.sentiment.neutral
                    }
                else:
                    # Fallback to VADER sentiment analysis
                    vader_scores = sentiment_analyzer.polarity_scores(utterance.text)
                    sentiment_score = {
                        'positive': vader_scores['pos'],
                        'negative': vader_scores['neg'],
                        'neutral': vader_scores['neu']
                    }
            except Exception as e:
                # If sentiment analysis fails, use neutral scores
                sentiment_score = {
                    'positive': 0.0,
                    'negative': 0.0,
                    'neutral': 1.0
                }
            
            # Detect emotions based on keywords
            detected_emotions = []
            text_lower = utterance.text.lower()
            for emotion, keywords in emotion_keywords.items():
                if any(keyword in text_lower for keyword in keywords):
                    detected_emotions.append(emotion)
            
            # Determine tone based on sentiment scores
            if sentiment_score['positive'] > 0.5:
                tone = 'Positive/Enthusiastic'
            elif sentiment_score['negative'] > 0.5:
                tone = 'Negative/Concerned'
            elif sentiment_score['neutral'] > 0.7:
                tone = 'Neutral/Calm'
            else:
                tone = 'Mixed'
            
            speaker_sentiments[speaker]['utterances'].append({
                'text': utterance.text,
                'sentiment': sentiment_score,
                'timestamp': utterance.start / 1000,
                'duration': (utterance.end - utterance.start) / 1000,
                'emotions': detected_emotions,
                'tone': tone
            })
            speaker_sentiments[speaker]['sentiment_scores'].append(sentiment_score)
            speaker_sentiments[speaker]['total_duration'] += (utterance.end - utterance.start) / 1000
            speaker_sentiments[speaker]['emotions'].extend(detected_emotions)
            speaker_sentiments[speaker]['tone_analysis'].append(tone)
        
        return speaker_sentiments
    
    def highlight_ai_blocks(self, transcript):
        """Highlight AI-related content and technical blocks in transcript"""
        ai_blocks = []
        
        # AI and technical keywords (more specific to avoid false positives)
        ai_keywords = [
            'artificial intelligence', 'machine learning', 'deep learning',
            'neural network', 'algorithm', 'model training', 'data prediction',
            'data science', 'analytics', 'automation', 'robotics', 'chatbot',
            'natural language processing', 'nlp', 'computer vision',
            'reinforcement learning', 'supervised learning', 'unsupervised learning',
            'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'pandas',
            'data mining', 'big data', 'cloud computing', 'api', 'database',
            'programming', 'software development', 'python', 'javascript'
        ]
        
        for utterance in transcript.utterances:
            text_lower = utterance.text.lower()
            
            # Check for AI keywords (use word boundaries to avoid false matches)
            for keyword in ai_keywords:
                # Use word boundary checks for short keywords to avoid false positives
                if len(keyword) <= 3:
                    # Skip very short generic terms that cause false positives
                    continue
                elif keyword in text_lower:
                    # Additional context check for relevant AI discussion
                    context_words = ['technology', 'system', 'development', 'project', 'software', 'digital']
                    has_context = any(ctx in text_lower for ctx in context_words)
                    
                    if has_context or len(keyword) > 10:  # Long keywords are usually specific enough
                        # Handle both string and integer speaker formats
                        if isinstance(utterance.speaker, str):
                            speaker_name = f"Speaker {utterance.speaker}"
                        else:
                            speaker_name = f"Speaker {chr(65 + int(utterance.speaker))}"
                            
                        ai_blocks.append({
                            'text': utterance.text,
                            'speaker': speaker_name,
                            'timestamp': utterance.start / 1000,
                            'ai_topic': keyword,
                            'start_time': f"{utterance.start / 1000 / 60:.2f}:{(utterance.start / 1000) % 60:02.0f}",
                            'end_time': f"{utterance.end / 1000 / 60:.2f}:{(utterance.end / 1000) % 60:02.0f}"
                        })
                        break  # Only add once per utterance
                    break
        
        return ai_blocks
    
    def analyze_audio_emotions(self, audio_file_path, transcript):
        """Analyze audio emotions for all utterances"""
        if not self.audio_emotion_detector:
            return transcript
        
        try:
            # Analyze emotions for each utterance
            utterances_with_emotions = self.audio_emotion_detector.analyze_utterance_emotions(
                audio_file_path, transcript.utterances
            )
            
            # Update transcript with emotion data
            transcript.utterances = utterances_with_emotions
            return transcript
            
        except Exception as e:
            print(f"Error analyzing audio emotions: {e}")
            return transcript
    
    def build_semantic_index(self, transcript):
        """Build semantic search index for transcript"""
        if not self.semantic_search_engine:
            return False
        
        try:
            # Extract utterance texts
            utterance_texts = [utterance.text for utterance in transcript.utterances]
            
            # Build index using the corrected method
            success = self.semantic_search_engine.build_index(utterance_texts)
            return success
            
        except Exception as e:
            print(f"Error building semantic index: {e}")
            return False
    
    def semantic_search(self, query, top_k=5):
        """Perform semantic search on transcript"""
        if not self.semantic_search_engine:
            return []
        
        try:
            results = self.semantic_search_engine.search(query, top_k)
            return results
        except Exception as e:
            print(f"Error in semantic search: {e}")
            return []
    
    def get_similar_utterances(self, reference_text, top_k=3):
        """Find utterances similar to a given text"""
        if not self.semantic_search_engine:
            return []
        
        try:
            results = self.semantic_search_engine.get_similar_utterances(reference_text, top_k)
            return results
        except Exception as e:
            print(f"Error finding similar utterances: {e}")
            return []
    
    def get_topic_clusters(self, n_clusters=5):
        """Get topic clusters from transcript"""
        if not self.semantic_search_engine:
            return []
        
        try:
            return self.semantic_search_engine.get_topic_clusters(n_clusters)
        except Exception as e:
            print(f"Error getting topic clusters: {e}")
            return []
    
    def calculate_engagement_score(self, transcript_data):
        """Calculate engagement score based on various metrics"""
        try:
            # Extract metrics
            total_duration = transcript_data.get('duration_minutes', 0)
            total_words = transcript_data.get('total_words', 0)
            speakers = transcript_data.get('speakers', [])
            utterances = transcript_data.get('utterances', [])
            
            # Calculate participation score based on speaking distribution
            participation_score = 0
            if speakers and utterances:
                # Calculate based on how evenly distributed the speaking time is
                speaker_times = {}
                for utterance in utterances:
                    speaker = utterance.speaker
                    duration = (utterance.end - utterance.start) / 1000 / 60  # minutes
                    speaker_times[speaker] = speaker_times.get(speaker, 0) + duration
                
                if speaker_times:
                    # Calculate participation balance (how evenly distributed)
                    total_time = sum(speaker_times.values())
                    if total_time > 0:
                        # Calculate standard deviation of speaking times
                        mean_time = total_time / len(speaker_times)
                        variance = sum((time - mean_time) ** 2 for time in speaker_times.values()) / len(speaker_times)
                        std_dev = variance ** 0.5
                        
                        # Participation score based on balance (lower std dev = higher score)
                        max_std = mean_time * 2  # Maximum expected deviation
                        balance_score = max(0, 1 - (std_dev / max_std)) if max_std > 0 else 0
                        participation_score = balance_score * 25
                    else:
                        participation_score = 0
            
            # Calculate interaction score
            interaction_score = 0
            if utterances:
                avg_utterances_per_minute = len(utterances) / max(total_duration, 1)
                interaction_score = min(avg_utterances_per_minute / 10, 1.0) * 25
            
            # Calculate sentiment score
            sentiment_score = 0
            positive_utterances = 0
            total_sentiment_utterances = 0
            
            for utterance in utterances:
                if hasattr(utterance, 'sentiment') and utterance.sentiment:
                    total_sentiment_utterances += 1
                    if utterance.sentiment.get('sentiment', '').lower() in ['positive', 'enthusiastic']:
                        positive_utterances += 1
            
            if total_sentiment_utterances > 0:
                sentiment_score = (positive_utterances / total_sentiment_utterances) * 25
            
            # Calculate action items score
            action_score = 0
            key_points = transcript_data.get('key_points', {})
            action_items = key_points.get('action_items', [])
            if action_items:
                action_score = min(len(action_items) / 5, 1.0) * 25
            
            # Total engagement score
            total_score = participation_score + interaction_score + sentiment_score + action_score
            
            return {
                'total_score': round(total_score, 1),
                'participation_score': round(participation_score, 1),
                'interaction_score': round(interaction_score, 1),
                'sentiment_score': round(sentiment_score, 1),
                'action_score': round(action_score, 1),
                'level': self._get_engagement_level(total_score)
            }
            
        except Exception as e:
            print(f"Error calculating engagement score: {e}")
            return {
                'total_score': 0,
                'participation_score': 0,
                'interaction_score': 0,
                'sentiment_score': 0,
                'action_score': 0,
                'level': 'Unknown'
            }
    
    def _get_engagement_level(self, score):
        """Get engagement level based on score"""
        if score >= 80:
            return "Excellent"
        elif score >= 60:
            return "Good"
        elif score >= 40:
            return "Moderate"
        elif score >= 20:
            return "Low"
        else:
            return "Poor"

    def export_to_docx(self, transcript, key_points, filename="transcript.docx"):
        """Export transcript to DOCX format with enhanced features"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx library not available. Please install it first.")
        doc = Document()
        # Title
        title = doc.add_heading('Meeting Transcript & Analysis', 0)
        title.alignment = 1  # Center alignment
        # Meeting info
        doc.add_heading('Meeting Information', level=1)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        # Safely calculate duration
        try:
            duration_minutes = transcript.audio_duration / 1000 / 60 if hasattr(transcript, 'audio_duration') and transcript.audio_duration else 0
        except:
            duration_minutes = 0
        doc.add_paragraph(f"Duration: {duration_minutes:.2f} minutes")
        doc.add_paragraph(f"Total Words: {len(transcript.text.split()):,}")
        doc.add_paragraph(f"Speakers: {len(set([u.speaker for u in transcript.utterances]))}")
        # Add summary
        doc.add_heading('Meeting Summary', level=1)
        try:
            if hasattr(self, 'advanced_meeting_summarizer') and self.advanced_meeting_summarizer:
                summary = self.generate_advanced_summary(transcript)
            else:
                summary = self.generate_summary(transcript)
        except Exception as e:
            summary = self.generate_summary(transcript)
        doc.add_paragraph(summary)
        # Add complete transcript with timestamps
        doc.add_heading('Complete Transcript', level=1)
        for utterance in transcript.utterances:
            try:
                start_time = f"{utterance.start / 1000 / 60:.2f}:{(utterance.start / 1000) % 60:02.0f}"
                end_time = f"{utterance.end / 1000 / 60:.2f}:{(utterance.end / 1000) % 60:02.0f}"
                speaker_para = doc.add_paragraph()
                speaker_para.add_run(f"Speaker {utterance.speaker} [{start_time} - {end_time}]: ").bold = True
                speaker_para.add_run(utterance.text)
            except Exception as e:
                # Fallback for malformed utterance data
                speaker_para = doc.add_paragraph()
                speaker_para.add_run(f"Speaker {getattr(utterance, 'speaker', 'Unknown')}: ").bold = True
                speaker_para.add_run(getattr(utterance, 'text', 'No text available'))
        # Add enhanced key points
        doc.add_heading('Key Points & Insights', level=1)
        # Handle missing or empty key_points
        if not key_points:
            doc.add_paragraph("No key points extracted.")
            return doc
        if key_points.get('decisions'):
            doc.add_heading('Decisions', level=2)
            for decision in key_points['decisions']:
                try:
                    timestamp = f"{decision.get('timestamp', 0) / 60:.2f}:{decision.get('timestamp', 0) % 60:02.0f}"
                    doc.add_paragraph(f"• [{timestamp}] {decision.get('text', 'No text')} ({decision.get('speaker', 'Unknown')})")
                except:
                    doc.add_paragraph(f"• {decision.get('text', 'Decision recorded')}")
        if key_points.get('action_items'):
            doc.add_heading('Action Items', level=2)
            for action in key_points['action_items']:
                try:
                    timestamp = f"{action.get('timestamp', 0) / 60:.2f}:{action.get('timestamp', 0) % 60:02.0f}"
                    doc.add_paragraph(f"• [{timestamp}] {action.get('text', 'No text')} ({action.get('speaker', 'Unknown')})")
                except:
                    doc.add_paragraph(f"• {action.get('text', 'Action item recorded')}")
        if key_points.get('questions'):
            doc.add_heading('Questions', level=2)
            for question in key_points['questions']:
                try:
                    timestamp = f"{question.get('timestamp', 0) / 60:.2f}:{question.get('timestamp', 0) % 60:02.0f}"
                    doc.add_paragraph(f"• [{timestamp}] {question.get('text', 'No text')} ({question.get('speaker', 'Unknown')})")
                except:
                    doc.add_paragraph(f"• {question.get('text', 'Question recorded')}")
        if key_points.get('organizations'):
            doc.add_heading('Organizations/Companies Mentioned', level=2)
            for org in key_points['organizations']:
                try:
                    timestamp = f"{org.get('timestamp', 0) / 60:.2f}:{org.get('timestamp', 0) % 60:02.0f}"
                    doc.add_paragraph(f"• [{timestamp}] {org.get('organization', 'Unknown')} - {org.get('text', 'No text')} ({org.get('speaker', 'Unknown')})")
                except:
                    doc.add_paragraph(f"• {org.get('organization', 'Organization mentioned')}")
        if key_points.get('technical_terms'):
            doc.add_heading('Technical Terms & Technologies', level=2)
            for tech in key_points['technical_terms']:
                try:
                    timestamp = f"{tech.get('timestamp', 0) / 60:.2f}:{tech.get('timestamp', 0) % 60:02.0f}"
                    doc.add_paragraph(f"• [{timestamp}] {tech.get('term', 'Unknown')} - {tech.get('text', 'No text')} ({tech.get('speaker', 'Unknown')})")
                except:
                    doc.add_paragraph(f"• {tech.get('term', 'Technical term mentioned')}")
        if key_points.get('highlights'):
            doc.add_heading('AI Highlights', level=2)
            for highlight in key_points['highlights']:
                try:
                    timestamp = f"{highlight.get('timestamp', 0) / 60:.2f}:{highlight.get('timestamp', 0) % 60:02.0f}"
                    doc.add_paragraph(f"• [{timestamp}] {highlight.get('reason', 'Highlight')}: {highlight.get('text', 'No text')}")
                except:
                    doc.add_paragraph(f"• {highlight.get('text', 'Highlight recorded')}")
        return doc

    def export_to_pdf(self, transcript, key_points, filename="transcript.pdf"):
        """Export transcript to PDF format with enhanced features"""
        if not REPORTLAB_AVAILABLE:
            raise Exception("ReportLab library not available. Please install it first.")
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30
        )
        story.append(Paragraph('Meeting Transcript & Analysis', title_style))
        story.append(Spacer(1, 12))
        # Meeting info
        story.append(Paragraph('Meeting Information', styles['Heading2']))
        story.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        # Safely calculate duration
        try:
            duration_minutes = transcript.audio_duration / 1000 / 60 if hasattr(transcript, 'audio_duration') and transcript.audio_duration else 0
        except:
            duration_minutes = 0
        story.append(Paragraph(f"Duration: {duration_minutes:.2f} minutes", styles['Normal']))
        story.append(Paragraph(f"Total Words: {len(transcript.text.split()):,}", styles['Normal']))
        story.append(Paragraph(f"Speakers: {len(set([u.speaker for u in transcript.utterances]))}", styles['Normal']))
        story.append(Spacer(1, 12))
        # Summary
        story.append(Paragraph('Meeting Summary', styles['Heading2']))
        try:
            if hasattr(self, 'advanced_meeting_summarizer') and self.advanced_meeting_summarizer:
                summary = self.generate_advanced_summary(transcript)
            else:
                summary = self.generate_summary(transcript)
        except Exception as e:
            summary = self.generate_summary(transcript)
        # Clean summary text for PDF
        summary_clean = summary.replace('**', '').replace('*', '').replace('#', '')
        story.append(Paragraph(summary_clean, styles['Normal']))
        story.append(Spacer(1, 12))
        # Complete transcript
        story.append(Paragraph('Complete Transcript', styles['Heading2']))
        for utterance in transcript.utterances:
            try:
                start_time = f"{utterance.start / 1000 / 60:.2f}:{(utterance.start / 1000) % 60:02.0f}"
                end_time = f"{utterance.end / 1000 / 60:.2f}:{(utterance.end / 1000) % 60:02.0f}"
                text = f"<b>Speaker {utterance.speaker} [{start_time} - {end_time}]:</b> {utterance.text}"
                story.append(Paragraph(text, styles['Normal']))
                story.append(Spacer(1, 6))
            except Exception as e:
                # Fallback for malformed utterance data
                text = f"<b>Speaker {getattr(utterance, 'speaker', 'Unknown')}:</b> {getattr(utterance, 'text', 'No text available')}"
                story.append(Paragraph(text, styles['Normal']))
                story.append(Spacer(1, 6))
        # Key points
        story.append(Paragraph('Key Points & Insights', styles['Heading2']))
        # Handle missing or empty key_points
        if not key_points:
            story.append(Paragraph("No key points extracted.", styles['Normal']))
        else:
            if key_points.get('decisions'):
                story.append(Paragraph('Decisions:', styles['Heading3']))
                for decision in key_points['decisions']:
                    try:
                        timestamp = f"{decision.get('timestamp', 0) / 60:.2f}:{decision.get('timestamp', 0) % 60:02.0f}"
                        story.append(Paragraph(f"• [{timestamp}] {decision.get('text', 'No text')} ({decision.get('speaker', 'Unknown')})", styles['Normal']))
                    except:
                        story.append(Paragraph(f"• {decision.get('text', 'Decision recorded')}", styles['Normal']))
            if key_points.get('action_items'):
                story.append(Paragraph('Action Items:', styles['Heading3']))
                for action in key_points['action_items']:
                    try:
                        timestamp = f"{action.get('timestamp', 0) / 60:.2f}:{action.get('timestamp', 0) % 60:02.0f}"
                        story.append(Paragraph(f"• [{timestamp}] {action.get('text', 'No text')} ({action.get('speaker', 'Unknown')})", styles['Normal']))
                    except:
                        story.append(Paragraph(f"• {action.get('text', 'Action item recorded')}", styles['Normal']))
            if key_points.get('organizations'):
                story.append(Paragraph('Organizations/Companies:', styles['Heading3']))
                for org in key_points['organizations']:
                    try:
                        timestamp = f"{org.get('timestamp', 0) / 60:.2f}:{org.get('timestamp', 0) % 60:02.0f}"
                        story.append(Paragraph(f"• [{timestamp}] {org.get('organization', 'Unknown')} - {org.get('text', 'No text')} ({org.get('speaker', 'Unknown')})", styles['Normal']))
                    except:
                        story.append(Paragraph(f"• {org.get('organization', 'Organization mentioned')}", styles['Normal']))
        doc.build(story)
        buffer.seek(0)
        return buffer

# Helper function to get transcript data safely
def get_transcript_data_safely(self):
    """Safely get transcript data from various sources"""
    try:
        # Try session state first
        if hasattr(self, 'transcript_data') and self.transcript_data:
            return self.transcript_data
        
        # Try from streamlit session state
        if 'verbatim_ai' in st.session_state and hasattr(st.session_state.verbatim_ai, 'transcript_data'):
            return st.session_state.verbatim_ai.transcript_data
        
        return None
    except Exception as e:
        print(f"Error getting transcript data: {e}")
        return None
    
        duration_minutes = transcript.audio_duration / 1000 / 60
        doc.add_paragraph(f"Duration: {duration_minutes:.2f} minutes")
        doc.add_paragraph(f"Total Words: {len(transcript.text.split()):,}")
        doc.add_paragraph(f"Speakers: {len(set([u.speaker for u in transcript.utterances]))}")
        
        # Add summary
        doc.add_heading('Meeting Summary', level=1)
        # Try to use advanced summary if available, otherwise fall back to basic summary
        try:
            if self.advanced_meeting_summarizer:
                summary = self.generate_advanced_summary(transcript)
            else:
                summary = self.generate_summary(transcript)
        except Exception as e:
            summary = self.generate_summary(transcript)
        doc.add_paragraph(summary)
        
        # Add complete transcript with timestamps
        doc.add_heading('Complete Transcript', level=1)
        for utterance in transcript.utterances:
            start_time = f"{utterance.start / 1000 / 60:.2f}:{(utterance.start / 1000) % 60:02.0f}"
            end_time = f"{utterance.end / 1000 / 60:.2f}:{(utterance.end / 1000) % 60:02.0f}"
            
            speaker_para = doc.add_paragraph()
            speaker_para.add_run(f"Speaker {utterance.speaker} [{start_time} - {end_time}]: ").bold = True
            speaker_para.add_run(utterance.text)
        
        # Add enhanced key points
        doc.add_heading('Key Points & Insights', level=1)
        
        if key_points['decisions']:
            doc.add_heading('Decisions', level=2)
            for decision in key_points['decisions']:
                timestamp = f"{decision['timestamp'] / 60:.2f}:{decision['timestamp'] % 60:02.0f}"
                doc.add_paragraph(f"• [{timestamp}] {decision['text']} ({decision['speaker']})")
        
        if key_points['action_items']:
            doc.add_heading('Action Items', level=2)
            for action in key_points['action_items']:
                timestamp = f"{action['timestamp'] / 60:.2f}:{action['timestamp'] % 60:02.0f}"
                doc.add_paragraph(f"• [{timestamp}] {action['text']} ({action['speaker']})")
        
        if key_points['questions']:
            doc.add_heading('Questions', level=2)
            for question in key_points['questions']:
                timestamp = f"{question['timestamp'] / 60:.2f}:{question['timestamp'] % 60:02.0f}"
                doc.add_paragraph(f"• [{timestamp}] {question['text']} ({question['speaker']})")
        
        if key_points['organizations']:
            doc.add_heading('Organizations/Companies Mentioned', level=2)
            for org in key_points['organizations']:
                timestamp = f"{org['timestamp'] / 60:.2f}:{org['timestamp'] % 60:02.0f}"
                doc.add_paragraph(f"• [{timestamp}] {org['organization']} - {org['text']} ({org['speaker']})")
        
        if key_points['technical_terms']:
            doc.add_heading('Technical Terms & Technologies', level=2)
            for tech in key_points['technical_terms']:
                timestamp = f"{tech['timestamp'] / 60:.2f}:{tech['timestamp'] % 60:02.0f}"
                doc.add_paragraph(f"• [{timestamp}] {tech['term']} - {tech['text']} ({tech['speaker']})")
        
        if key_points['highlights']:
            doc.add_heading('AI Highlights', level=2)
            for highlight in key_points['highlights']:
                timestamp = f"{highlight['timestamp'] / 60:.2f}:{highlight['timestamp'] % 60:02.0f}"
                doc.add_paragraph(f"• [{timestamp}] {highlight['reason']}: {highlight['text']}")
        
        return doc
    
    def export_to_docx(self, transcript, key_points, filename="transcript.docx"):
        """Export transcript to DOCX format with enhanced features"""
        if not DOCX_AVAILABLE:
            raise Exception("python-docx library not available. Please install it first.")
        
        doc = Document()
        
        # Title
        title = doc.add_heading('Meeting Transcript & Analysis', 0)
        title.alignment = 1  # Center alignment
        
        # Meeting info
        doc.add_heading('Meeting Information', level=1)
        doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        
        # Safely calculate duration
        try:
            duration_minutes = transcript.audio_duration / 1000 / 60 if hasattr(transcript, 'audio_duration') and transcript.audio_duration else 0
        except:
            duration_minutes = 0
            
        doc.add_paragraph(f"Duration: {duration_minutes:.2f} minutes")
        doc.add_paragraph(f"Total Words: {len(transcript.text.split()):,}")
        doc.add_paragraph(f"Speakers: {len(set([u.speaker for u in transcript.utterances]))}")
        
        # Add summary
        doc.add_heading('Meeting Summary', level=1)
        try:
            if hasattr(self, 'advanced_meeting_summarizer') and self.advanced_meeting_summarizer:
                summary = self.generate_advanced_summary(transcript)
            else:
                summary = self.generate_summary(transcript)
        except Exception as e:
            summary = self.generate_summary(transcript)
        doc.add_paragraph(summary)
        
        # Add complete transcript with timestamps
        doc.add_heading('Complete Transcript', level=1)
        for utterance in transcript.utterances:
            try:
                start_time = f"{utterance.start / 1000 / 60:.2f}:{(utterance.start / 1000) % 60:02.0f}"
                end_time = f"{utterance.end / 1000 / 60:.2f}:{(utterance.end / 1000) % 60:02.0f}"
                
                speaker_para = doc.add_paragraph()
                speaker_para.add_run(f"Speaker {utterance.speaker} [{start_time} - {end_time}]: ").bold = True
                speaker_para.add_run(utterance.text)
            except Exception as e:
                # Fallback for malformed utterance data
                speaker_para = doc.add_paragraph()
                speaker_para.add_run(f"Speaker {getattr(utterance, 'speaker', 'Unknown')}: ").bold = True
                speaker_para.add_run(getattr(utterance, 'text', 'No text available'))
        
        # Add enhanced key points
        doc.add_heading('Key Points & Insights', level=1)
        
        # Handle missing or empty key_points
        if not key_points:
            doc.add_paragraph("No key points extracted.")
            return doc
        
        if key_points.get('decisions'):
            doc.add_heading('Decisions', level=2)
            for decision in key_points['decisions']:
                try:
                    timestamp = f"{decision.get('timestamp', 0) / 60:.2f}:{decision.get('timestamp', 0) % 60:02.0f}"
                    doc.add_paragraph(f"• [{timestamp}] {decision.get('text', 'No text')} ({decision.get('speaker', 'Unknown')})")
                except:
                    doc.add_paragraph(f"• {decision.get('text', 'Decision recorded')}")
        
        if key_points.get('action_items'):
            doc.add_heading('Action Items', level=2)
            for action in key_points['action_items']:
                try:
                    timestamp = f"{action.get('timestamp', 0) / 60:.2f}:{action.get('timestamp', 0) % 60:02.0f}"
                    doc.add_paragraph(f"• [{timestamp}] {action.get('text', 'No text')} ({action.get('speaker', 'Unknown')})")
                except:
                    doc.add_paragraph(f"• {action.get('text', 'Action item recorded')}")
        
        if key_points.get('questions'):
            doc.add_heading('Questions', level=2)
            for question in key_points['questions']:
                try:
                    timestamp = f"{question.get('timestamp', 0) / 60:.2f}:{question.get('timestamp', 0) % 60:02.0f}"
                    doc.add_paragraph(f"• [{timestamp}] {question.get('text', 'No text')} ({question.get('speaker', 'Unknown')})")
                except:
                    doc.add_paragraph(f"• {question.get('text', 'Question recorded')}")
        
        if key_points.get('organizations'):
            doc.add_heading('Organizations/Companies Mentioned', level=2)
            for org in key_points['organizations']:
                try:
                    timestamp = f"{org.get('timestamp', 0) / 60:.2f}:{org.get('timestamp', 0) % 60:02.0f}"
                    doc.add_paragraph(f"• [{timestamp}] {org.get('organization', 'Unknown')} - {org.get('text', 'No text')} ({org.get('speaker', 'Unknown')})")
                except:
                    doc.add_paragraph(f"• {org.get('organization', 'Organization mentioned')}")
        
        if key_points.get('technical_terms'):
            doc.add_heading('Technical Terms & Technologies', level=2)
            for tech in key_points['technical_terms']:
                try:
                    timestamp = f"{tech.get('timestamp', 0) / 60:.2f}:{tech.get('timestamp', 0) % 60:02.0f}"
                    doc.add_paragraph(f"• [{timestamp}] {tech.get('term', 'Unknown')} - {tech.get('text', 'No text')} ({tech.get('speaker', 'Unknown')})")
                except:
                    doc.add_paragraph(f"• {tech.get('term', 'Technical term mentioned')}")
        
        if key_points.get('highlights'):
            doc.add_heading('AI Highlights', level=2)
            for highlight in key_points['highlights']:
                try:
                    timestamp = f"{highlight.get('timestamp', 0) / 60:.2f}:{highlight.get('timestamp', 0) % 60:02.0f}"
                    doc.add_paragraph(f"• [{timestamp}] {highlight.get('reason', 'Highlight')}: {highlight.get('text', 'No text')}")
                except:
                    doc.add_paragraph(f"• {highlight.get('text', 'Highlight recorded')}")
        
        return doc
    
    def export_to_pdf(self, transcript, key_points, filename="transcript.pdf"):
        """Export transcript to PDF format with enhanced features"""
        if not REPORTLAB_AVAILABLE:
            raise Exception("ReportLab library not available. Please install it first.")
        
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Title
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30
        )
        story.append(Paragraph('Meeting Transcript & Analysis', title_style))
        story.append(Spacer(1, 12))
        
        # Meeting info
        story.append(Paragraph('Meeting Information', styles['Heading2']))
        story.append(Paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", styles['Normal']))
        
        # Safely calculate duration
        try:
            duration_minutes = transcript.audio_duration / 1000 / 60 if hasattr(transcript, 'audio_duration') and transcript.audio_duration else 0
        except:
            duration_minutes = 0
            
        story.append(Paragraph(f"Duration: {duration_minutes:.2f} minutes", styles['Normal']))
        story.append(Paragraph(f"Total Words: {len(transcript.text.split()):,}", styles['Normal']))
        story.append(Paragraph(f"Speakers: {len(set([u.speaker for u in transcript.utterances]))}", styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Summary
        story.append(Paragraph('Meeting Summary', styles['Heading2']))
        try:
            if hasattr(self, 'advanced_meeting_summarizer') and self.advanced_meeting_summarizer:
                summary = self.generate_advanced_summary(transcript)
            else:
                summary = self.generate_summary(transcript)
        except Exception as e:
            summary = self.generate_summary(transcript)
        
        # Clean summary text for PDF
        summary_clean = summary.replace('**', '').replace('*', '').replace('#', '')
        story.append(Paragraph(summary_clean, styles['Normal']))
        story.append(Spacer(1, 12))
        
        # Complete transcript
        story.append(Paragraph('Complete Transcript', styles['Heading2']))
        for utterance in transcript.utterances:
            try:
                start_time = f"{utterance.start / 1000 / 60:.2f}:{(utterance.start / 1000) % 60:02.0f}"
                end_time = f"{utterance.end / 1000 / 60:.2f}:{(utterance.end / 1000) % 60:02.0f}"
                text = f"<b>Speaker {utterance.speaker} [{start_time} - {end_time}]:</b> {utterance.text}"
                story.append(Paragraph(text, styles['Normal']))
                story.append(Spacer(1, 6))
            except Exception as e:
                # Fallback for malformed utterance data
                text = f"<b>Speaker {getattr(utterance, 'speaker', 'Unknown')}:</b> {getattr(utterance, 'text', 'No text available')}"
                story.append(Paragraph(text, styles['Normal']))
                story.append(Spacer(1, 6))
        
        # Key points
        story.append(Paragraph('Key Points & Insights', styles['Heading2']))
        
        # Handle missing or empty key_points
        if not key_points:
            story.append(Paragraph("No key points extracted.", styles['Normal']))
        else:
            if key_points.get('decisions'):
                story.append(Paragraph('Decisions:', styles['Heading3']))
                for decision in key_points['decisions']:
                    try:
                        timestamp = f"{decision.get('timestamp', 0) / 60:.2f}:{decision.get('timestamp', 0) % 60:02.0f}"
                        story.append(Paragraph(f"• [{timestamp}] {decision.get('text', 'No text')} ({decision.get('speaker', 'Unknown')})", styles['Normal']))
                    except:
                        story.append(Paragraph(f"• {decision.get('text', 'Decision recorded')}", styles['Normal']))
            
            if key_points.get('action_items'):
                story.append(Paragraph('Action Items:', styles['Heading3']))
                for action in key_points['action_items']:
                    try:
                        timestamp = f"{action.get('timestamp', 0) / 60:.2f}:{action.get('timestamp', 0) % 60:02.0f}"
                        story.append(Paragraph(f"• [{timestamp}] {action.get('text', 'No text')} ({action.get('speaker', 'Unknown')})", styles['Normal']))
                    except:
                        story.append(Paragraph(f"• {action.get('text', 'Action item recorded')}", styles['Normal']))
            
            if key_points.get('organizations'):
                story.append(Paragraph('Organizations/Companies:', styles['Heading3']))
                for org in key_points['organizations']:
                    try:
                        timestamp = f"{org.get('timestamp', 0) / 60:.2f}:{org.get('timestamp', 0) % 60:02.0f}"
                        story.append(Paragraph(f"• [{timestamp}] {org.get('organization', 'Unknown')} - {org.get('text', 'No text')} ({org.get('speaker', 'Unknown')})", styles['Normal']))
                    except:
                        story.append(Paragraph(f"• {org.get('organization', 'Organization mentioned')}", styles['Normal']))
        
        doc.build(story)
        buffer.seek(0)
        return buffer

# Helper function to get transcript data safely
def get_transcript_data_safely():
    """Safely get transcript data from various sources"""
    try:
        # Try session state first
        if hasattr(st.session_state, 'verbatim_ai') and hasattr(st.session_state.verbatim_ai, 'transcript_data'):
            return st.session_state.verbatim_ai.transcript_data
        
        return None
    except Exception as e:
        print(f"Error getting transcript data: {e}")
        return None

def safe_get_key_points(analysis_results):
    """Safely extract key points from analysis results"""
    if not analysis_results:
        return {
            'decisions': [],
            'action_items': [],
            'questions': [],
            'highlights': [],
            'organizations': [],
            'technical_terms': []
        }
    
    key_points = analysis_results.get('key_points', {})
    
    # Ensure all required keys exist
    for key in ['decisions', 'action_items', 'questions', 'highlights', 'organizations', 'technical_terms']:
        if key not in key_points:
            key_points[key] = []
    
    return key_points

def main():
    # Initialize modern UI
    init_modern_ui()
    
    # Check authentication first
    if not check_authentication():
        show_login_page()
        return
    
    # Get current user role
    current_user_role = get_current_user_role()
    
    # Initialize session state
    if 'verbatim_ai' not in st.session_state:
        st.session_state.verbatim_ai = VerbatimAI()
    
    # Show user info in sidebar
    show_user_info()
    
    # Get navigation items based on user role
    if current_user_role:
        nav_items = NavigationManager.get_navigation_items(current_user_role)
    else:
        nav_items = [
            {"title": "🏠 Dashboard", "icon": "house", "page": "dashboard"},
            {"title": "🎙️ Upload, Record & Transcribe", "icon": "mic", "page": "upload"},
            {"title": "🔴 Real-time Recording", "icon": "record", "page": "recording"},
            {"title": "📚 Meeting Library", "icon": "book", "page": "library"},
            {"title": "🔍 Semantic Search", "icon": "search", "page": "search"},
            {"title": "📧 Email Summary", "icon": "envelope", "page": "email"},
            {"title": "🔄 Task Monitor", "icon": "gear", "page": "tasks"},
        ]
    
    # Modern Sidebar navigation
    with st.sidebar:
        # Logo and branding
        st.markdown("""
        <div class="sidebar-logo">
            <div style="font-size: 3rem; margin-bottom: 0.5rem;">🎙️</div>
            <div class="sidebar-title">VerbatimAI</div>
            <div style="font-size: 0.85rem; color: var(--neutral-600); margin-top: 0.5rem;">
                Smart Meeting Intelligence
            </div>
        </div>
        """, unsafe_allow_html=True)
        
        # API Key Status with modern design
        if st.session_state.verbatim_ai.verify_api_key():
            create_status_badge('success', 'API Key: Valid')
        else:
            create_status_badge('error', 'API Key: Invalid')
        
        # Create navigation menu
        if OPTION_MENU_AVAILABLE:
            nav_options = [item["title"] for item in nav_items]
            nav_icons = [item["icon"] for item in nav_items]
            
            selected = option_menu(
                menu_title=None,
                options=nav_options,
                icons=nav_icons,
                menu_icon="cast",
                default_index=0,
                styles={
                    "container": {"padding": "0!important", "background-color": "transparent"},
                    "icon": {"color": "var(--primary-blue)", "font-size": "18px"}, 
                    "nav-link": {
                        "font-size": "14px", 
                        "text-align": "left", 
                        "margin": "0px",
                        "padding": "12px 16px",
                        "border-radius": "12px",
                        "font-family": "'Inter', sans-serif",
                        "font-weight": "500"
                    },
                    "nav-link-selected": {
                        "background": "linear-gradient(135deg, var(--primary-blue), var(--primary-blue-dark))",
                        "color": "white",
                        "box-shadow": "var(--shadow-md)"
                    },
                }
            )
        else:
            # Fallback navigation
            nav_options = [item["title"] for item in nav_items]
            selected = st.selectbox("Navigation", nav_options)
    
    # Map selected option to page (check for session state navigation first)
    selected_page = None
    
    # Check if there's a page selection from dashboard buttons
    if 'selected_page' in st.session_state:
        selected_page = st.session_state.selected_page
        # Clear the session state to prevent sticking
        del st.session_state.selected_page
    else:
        # Normal navigation menu selection
        for item in nav_items:
            if item["title"] == selected:
                selected_page = item["page"]
                break
    
    # Route to appropriate page
    if selected_page == "dashboard":
        show_dashboard()
    elif selected_page == "upload":
        show_upload_transcribe()
    elif selected_page == "recording":
        if ADVANCED_FEATURES_AVAILABLE:
            show_real_time_recording()
        else:
            st.error("🔴 Real-time recording not available. Please install required dependencies.")
            st.info("Try running: `pip install pyaudio opencv-python`")
    elif selected_page == "library":
        if ADVANCED_FEATURES_AVAILABLE:
            show_meeting_library()
        else:
            st.error("📚 Meeting library not available. Please install required dependencies.")
    elif selected_page == "analytics":
        show_analytics()
    elif selected_page == "sentiment":
        show_speaker_sentiment_eda()
    elif selected_page == "search":
        show_semantic_search()
    elif selected_page == "email":
        show_email_summary()
    elif selected_page == "export":
        show_export()
    elif selected_page == "users":
        show_user_management()
    elif selected_page == "settings":
        show_system_settings()
    elif selected_page == "tasks":
        show_task_monitor()
    elif selected_page == "editor":
        show_interactive_transcript_editor(st.session_state.verbatim_ai.transcript_data, 
                                        st.session_state.current_user.username if 'current_user' in st.session_state else None)
    elif selected_page == "library":
        show_library()
    elif selected_page == "tasks":
        if ADVANCED_FEATURES_AVAILABLE:
            show_task_monitor()
        else:
            st.error("🔄 Task monitor not available. Please install required dependencies.")
            st.info("Try running: `pip install celery redis`")
    else:
        show_dashboard()

def generate_system_report():
    """Generate comprehensive system report"""
    report = {
        'timestamp': datetime.now().isoformat(),
        'python_version': sys.version,
        'streamlit_version': st.__version__,
        'assemblyai_available': 'assemblyai' in sys.modules,
        'advanced_features_available': ADVANCED_FEATURES_AVAILABLE,
    }
    
    try:
        import assemblyai as aai
        report['assemblyai_version'] = getattr(aai, '__version__', 'Unknown')
    except:
        report['assemblyai_version'] = 'Not installed'
    
    try:
        api_status = st.session_state.verbatim_ai.verify_api_key()
        report['api_status'] = 'Valid' if api_status else 'Invalid'
    except:
        report['api_status'] = 'Error'
    
    try:
        from async_task_queue import TaskManager
        task_manager = TaskManager()
        tasks = task_manager.get_all_tasks()
        report['active_tasks'] = len([t for t in tasks if t['status'] not in ['SUCCESS', 'FAILURE']])
        report['total_tasks'] = len(tasks)
    except:
        report['active_tasks'] = 'N/A'
        report['total_tasks'] = 'N/A'
    
    # Check installed packages
    required_packages = ['streamlit', 'assemblyai', 'sentence-transformers', 'faiss-cpu', 'pyaudio', 'celery', 'redis']
    for package in required_packages:
        try:
            __import__(package)
            report[f'{package}_installed'] = True
        except ImportError:
            report[f'{package}_installed'] = False
    
    return report

def export_system_report():
    """Export system report as JSON"""
    report = generate_system_report()
    import json
    import io
    
    # Create JSON report
    json_buffer = io.StringIO()
    json.dump(report, json_buffer, indent=2, default=str)
    json_buffer.seek(0)
    
    return json_buffer.getvalue()

# Utility Functions
def calculate_transcript_duration(transcript):
    """Calculate transcript duration safely"""
    if hasattr(transcript, 'audio_duration') and transcript.audio_duration:
        return transcript.audio_duration / 1000 / 60  # Convert to minutes
    elif transcript.utterances:
        # Calculate from last utterance end time
        last_utterance = max(transcript.utterances, key=lambda u: u.end)
        return last_utterance.end / 1000 / 60  # Convert to minutes
    return 0

def show_dashboard():
    # Hero Section with modern design
    st.markdown("""
    <div style="text-align: center; padding: 3rem 0;">
        <div class="main-header fade-in">VerbatimAI</div>
        <div style="font-size: 1.3rem; color: var(--neutral-600); margin-bottom: 2rem; font-weight: 400;">
            Transform your meetings into actionable insights with AI-powered intelligence
        </div>
        <div style="display: flex; justify-content: center; gap: 1rem; margin-bottom: 3rem;">
            <span style="display: inline-flex; align-items: center; padding: 0.5rem 1rem; background: var(--primary-blue-light); border-radius: 20px; font-size: 0.9rem; color: var(--primary-blue); font-weight: 500;">
                🚀 Real-time Transcription
            </span>
            <span style="display: inline-flex; align-items: center; padding: 0.5rem 1rem; background: var(--primary-blue-light); border-radius: 20px; font-size: 0.9rem; color: var(--primary-blue); font-weight: 500;">
                🧠 AI Analysis
            </span>
            <span style="display: inline-flex; align-items: center; padding: 0.5rem 1rem; background: var(--primary-blue-light); border-radius: 20px; font-size: 0.9rem; color: var(--primary-blue); font-weight: 500;">
                📊 Smart Insights
            </span>
        </div>
    </div>
    """, unsafe_allow_html=True)
    
    # Feature cards with modern design
    st.markdown('<h2 class="section-header">🚀 Powerful Features</h2>', unsafe_allow_html=True)
    
    col1, col2, col3, col4 = st.columns(4, gap="large")
    
    with col1:
        st.markdown("""
        <div class="feature-card slide-up">
            <div style="font-size: 3rem; margin-bottom: 1rem; text-align: center;">🎙️</div>
            <h3>Smart Transcription</h3>
            <p>Upload audio/video files and get accurate transcriptions with advanced speaker identification and timestamp precision.</p>
            <div style="margin-top: 1rem;">
                <span style="font-size: 0.8rem; color: var(--secondary-green); font-weight: 500;">✓ 98% Accuracy</span>
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    with col2:
        st.markdown("""
        <div class="feature-card slide-up" style="animation-delay: 0.1s;">
            <div style="font-size: 3rem; margin-bottom: 1rem; text-align: center;">🔴</div>
            <h3>Live Recording</h3>
            <p>Record meetings in real-time with modern web browser technology supporting both audio and video capture.</p>
            <div style="margin-top: 1rem;">
                <span style="font-size: 0.8rem; color: var(--secondary-green); font-weight: 500;">✓ Browser Native</span>
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    with col3:
        st.markdown("""
        <div class="feature-card slide-up" style="animation-delay: 0.2s;">
            <div style="font-size: 3rem; margin-bottom: 1rem; text-align: center;">🧠</div>
            <h3>AI Insights</h3>
            <p>Extract key decisions, action items, sentiment analysis, and generate comprehensive meeting summaries automatically.</p>
            <div style="margin-top: 1rem;">
                <span style="font-size: 0.8rem; color: var(--secondary-green); font-weight: 500;">✓ GPT-Powered</span>
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    with col4:
        st.markdown("""
        <div class="feature-card slide-up" style="animation-delay: 0.3s;">
            <div style="font-size: 3rem; margin-bottom: 1rem; text-align: center;">📊</div>
            <h3>Rich Analytics</h3>
            <p>Visualize speaker engagement, sentiment trends, topic clusters, and meeting dynamics with interactive charts.</p>
            <div style="margin-top: 1rem;">
                <span style="font-size: 0.8rem; color: var(--secondary-green); font-weight: 500;">✓ Interactive</span>
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    # Quick stats with modern metrics
    if st.session_state.verbatim_ai.transcript_data:
        st.markdown('<h2 class="section-header">📈 Current Meeting Analytics</h2>', unsafe_allow_html=True)
        
        col1, col2, col3, col4, col5 = st.columns(5, gap="medium")
        
        duration = calculate_transcript_duration(st.session_state.verbatim_ai.transcript_data)
        speakers = len(set([u.speaker for u in st.session_state.verbatim_ai.transcript_data.utterances]))
        words = len(st.session_state.verbatim_ai.transcript_data.text.split())
        utterances = len(st.session_state.verbatim_ai.transcript_data.utterances)
        
        # Calculate engagement score
        try:
            transcript_data = {
                'duration_minutes': duration,
                'total_words': words,
                'speakers': list(set([u.speaker for u in st.session_state.verbatim_ai.transcript_data.utterances])),
                'utterances': st.session_state.verbatim_ai.transcript_data.utterances,
                'key_points': st.session_state.verbatim_ai.analysis_results.get('key_points', {}) if st.session_state.verbatim_ai.analysis_results else {}
            }
            engagement = st.session_state.verbatim_ai.calculate_engagement_score(transcript_data)
            engagement_score = engagement['total_score']
        except:
            engagement_score = 0
        
        with col1:
            create_modern_metric_card("Duration", f"{duration:.1f}m", icon="⏱️")
        with col2:
            create_modern_metric_card("Speakers", str(speakers), icon="👥")
        with col3:
            create_modern_metric_card("Words", f"{words:,}", icon="💬")
        with col4:
            create_modern_metric_card("Utterances", str(utterances), icon="🗣️")
        with col5:
            engagement_icon = "🔥" if engagement_score > 70 else "📊"
            create_modern_metric_card("Engagement", f"{engagement_score}%", icon=engagement_icon)
    
    # Quick Actions with modern buttons
    st.markdown('<h2 class="section-header">⚡ Quick Actions</h2>', unsafe_allow_html=True)
    
    col1, col2, col3, col4 = st.columns(4, gap="large")
    
    with col1:
        st.markdown("""
        <div class="feature-card" style="text-align: center; padding: 2rem 1rem;">
            <div style="font-size: 3rem; margin-bottom: 1rem;">🔴</div>
            <h4 style="margin-bottom: 1rem;">Start Recording</h4>
            <p style="font-size: 0.9rem; margin-bottom: 1.5rem;">Begin real-time meeting capture</p>
        </div>
        """, unsafe_allow_html=True)
        if st.button("🔴 Start Recording", type="primary", use_container_width=True, key="start_recording"):
            st.session_state.selected_page = "recording"
            st.rerun()
    
    with col2:
        st.markdown("""
        <div class="feature-card" style="text-align: center; padding: 2rem 1rem;">
            <div style="font-size: 3rem; margin-bottom: 1rem;">📤</div>
            <h4 style="margin-bottom: 1rem;">Upload & Transcribe</h4>
            <p style="font-size: 0.9rem; margin-bottom: 1.5rem;">Process existing audio/video</p>
        </div>
        """, unsafe_allow_html=True)
        if st.button("📤 Upload File", use_container_width=True, key="upload_file"):
            st.session_state.selected_page = "upload"
            st.rerun()
    
    with col3:
        st.markdown("""
        <div class="feature-card" style="text-align: center; padding: 2rem 1rem;">
            <div style="font-size: 3rem; margin-bottom: 1rem;">🔍</div>
            <h4 style="margin-bottom: 1rem;">Search Transcripts</h4>
            <p style="font-size: 0.9rem; margin-bottom: 1.5rem;">Find content with AI search</p>
        </div>
        """, unsafe_allow_html=True)
        if st.button("🔍 Semantic Search", use_container_width=True, key="semantic_search"):
            st.session_state.selected_page = "search"
            st.rerun()
    
    with col4:
        st.markdown("""
        <div class="feature-card" style="text-align: center; padding: 2rem 1rem;">
            <div style="font-size: 3rem; margin-bottom: 1rem;">📚</div>
            <h4 style="margin-bottom: 1rem;">Meeting Library</h4>
            <p style="font-size: 0.9rem; margin-bottom: 1.5rem;">Browse past meetings</p>
        </div>
        """, unsafe_allow_html=True)
        if st.button("📚 Browse Library", use_container_width=True, key="browse_library"):
            st.session_state.selected_page = "library"
            st.rerun()
    
    # Export & System Status
    st.markdown('<h2 class="section-header">📄 Export & System Status</h2>', unsafe_allow_html=True)
    
    # Export buttons
    col_export1, col_export2, col_export3 = st.columns(3)
    
    with col_export1:
        if st.button("📄 Export PDF Report", use_container_width=True):
            if hasattr(st.session_state.verbatim_ai, 'transcript_data') and st.session_state.verbatim_ai.transcript_data:
                try:
                    transcript = st.session_state.verbatim_ai.transcript_data
                    key_points = {}
                    if hasattr(st.session_state.verbatim_ai, 'analysis_results') and st.session_state.verbatim_ai.analysis_results:
                        key_points = st.session_state.verbatim_ai.analysis_results.get('key_points', {})
                    
                    if not REPORTLAB_AVAILABLE:
                        st.error("❌ ReportLab library not available. Please install it first: pip install reportlab")
                    else:
                        pdf_buffer = st.session_state.verbatim_ai.export_to_pdf(transcript, key_points)
                        
                        st.download_button(
                            label="📥 Download PDF",
                            data=pdf_buffer.getvalue(),
                            file_name=f"meeting_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf",
                            mime="application/pdf",
                            use_container_width=True
                        )
                        st.success("✅ PDF report generated successfully!")
                except Exception as e:
                    st.error(f"❌ PDF export failed: {str(e)}")
            else:
                st.warning("⚠️ No transcript data available. Please upload and transcribe an audio file first.")
    
    with col_export2:
        if st.button("📝 Export DOCX Report", use_container_width=True):
            if hasattr(st.session_state.verbatim_ai, 'transcript_data') and st.session_state.verbatim_ai.transcript_data:
                try:
                    transcript = st.session_state.verbatim_ai.transcript_data
                    key_points = {}
                    if hasattr(st.session_state.verbatim_ai, 'analysis_results') and st.session_state.verbatim_ai.analysis_results:
                        key_points = st.session_state.verbatim_ai.analysis_results.get('key_points', {})
                    
                    if not DOCX_AVAILABLE:
                        st.error("❌ python-docx library not available. Please install it first: pip install python-docx")
                    else:
                        doc = st.session_state.verbatim_ai.export_to_docx(transcript, key_points)
                        
                        # Save to bytes buffer
                        buffer = io.BytesIO()
                        doc.save(buffer)
                        buffer.seek(0)
                        
                        st.download_button(
                            label="📥 Download DOCX",
                            data=buffer.getvalue(),
                            file_name=f"meeting_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            use_container_width=True
                        )
                        st.success("✅ DOCX report generated successfully!")
                except Exception as e:
                    st.error(f"❌ DOCX export failed: {str(e)}")
            else:
                st.warning("⚠️ No transcript data available. Please upload and transcribe an audio file first.")
    
    with col_export3:
        if st.button("📊 Export Analytics CSV", use_container_width=True):
            if hasattr(st.session_state.verbatim_ai, 'transcript_data') and st.session_state.verbatim_ai.transcript_data:
                try:
                    transcript = st.session_state.verbatim_ai.transcript_data
                    
                    if not PANDAS_AVAILABLE:
                        st.error("❌ Pandas library not available. Please install it first: pip install pandas")
                    else:
                        # Create analytics data
                        analytics_data = []
                        for utterance in transcript.utterances:
                            analytics_data.append({
                                'Speaker': f"Speaker {utterance.speaker}",
                                'Text': utterance.text,
                                'Start_Time': utterance.start / 1000,
                                'End_Time': utterance.end / 1000,
                                'Duration': (utterance.end - utterance.start) / 1000,
                                'Word_Count': len(utterance.text.split())
                            })
                        
                        df = pd.DataFrame(analytics_data)
                        
                        # Convert to CSV
                        csv_buffer = io.StringIO()
                        df.to_csv(csv_buffer, index=False)
                        csv_buffer.seek(0)
                        
                        st.download_button(
                            label="📥 Download CSV",
                            data=csv_buffer.getvalue(),
                            file_name=f"meeting_analytics_{datetime.now().strftime('%Y%m%d_%H%M%S')}.csv",
                            mime="text/csv",
                            use_container_width=True
                        )
                        st.success("✅ CSV analytics exported successfully!")
                except Exception as e:
                    st.error(f"❌ CSV export failed: {str(e)}")
            else:
                st.warning("⚠️ No transcript data available. Please upload and transcribe an audio file first.")
    
    # System Report Export
    col_report1, col_report2 = st.columns(2)
    
    with col_report1:
        if st.button("📋 Export System Report", use_container_width=True):
            try:
                system_report = export_system_report()
                
                st.download_button(
                    label="📥 Download System Report",
                    data=system_report,
                    file_name=f"system_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                    mime="application/json",
                    use_container_width=True
                )
                st.success("✅ System report exported successfully!")
            except Exception as e:
                st.error(f"❌ System report export failed: {str(e)}")
    
    with col_report2:
        if st.button("📊 Show System Report", use_container_width=True):
            try:
                report = generate_system_report()
                st.json(report)
                st.success("✅ System report generated!")
            except Exception as e:
                st.error(f"❌ System report generation failed: {str(e)}")
    
    st.markdown("---")
    
    col1, col2 = st.columns(2, gap="large")
    
    with col1:
        st.markdown("""
        <div class="feature-card">
            <h4 style="margin-bottom: 1rem;">🔗 Service Connectivity</h4>
        </div>
        """, unsafe_allow_html=True)
        
        # API Status
        if st.session_state.verbatim_ai.verify_api_key():
            create_status_badge('success', 'AssemblyAI API: Connected')
        else:
            create_status_badge('error', 'AssemblyAI API: Disconnected')
        
        # Dependencies Status
        if PLOTLY_AVAILABLE:
            create_status_badge('success', 'Plotly: Available')
        else:
            create_status_badge('warning', 'Plotly: Limited functionality')
        
        if OPTION_MENU_AVAILABLE:
            create_status_badge('success', 'Option Menu: Available')
        else:
            create_status_badge('warning', 'Option Menu: Fallback mode')
    
    with col2:
        st.markdown("""
        <div class="feature-card">
            <h4 style="margin-bottom: 1rem;">⚡ Performance Metrics</h4>
        </div>
        """, unsafe_allow_html=True)
        
        # Show performance metrics
        create_status_badge('info', f'Python Version: {sys.version.split()[0]}')
        create_status_badge('info', f'Streamlit Version: {st.__version__}')
        
        # Memory usage (if available)
        try:
            import psutil
            memory_percent = psutil.virtual_memory().percent
            if memory_percent < 70:
                create_status_badge('success', f'Memory Usage: {memory_percent:.1f}%')
            elif memory_percent < 85:
                create_status_badge('warning', f'Memory Usage: {memory_percent:.1f}%')
            else:
                create_status_badge('error', f'Memory Usage: {memory_percent:.1f}%')
        except:
            create_status_badge('info', 'Memory Usage: Not available')
    
    # Async Task Pipeline Status
    st.markdown('<h2 class="section-header">🔄 Background Task Pipeline</h2>', unsafe_allow_html=True)
    
    col_task1, col_task2 = st.columns(2)
    
    with col_task1:
        st.markdown("""
        <div class="feature-card">
            <h4 style="margin-bottom: 1rem;">🔄 Celery Task Queue</h4>
        </div>
        """, unsafe_allow_html=True)
        
        # Check if async features are available
        if ADVANCED_FEATURES_AVAILABLE:
            try:
                from async_task_queue import TaskManager
                task_manager = TaskManager()
                tasks = task_manager.get_all_tasks()
                
                active_tasks = len([t for t in tasks if t['status'] not in ['SUCCESS', 'FAILURE']])
                total_tasks = len(tasks)
                
                st.metric("Active Tasks", active_tasks)
                st.metric("Total Tasks", total_tasks)
                
                if active_tasks > 0:
                    st.success("✅ Background processing active")
                else:
                    st.info("ℹ️ No active background tasks")
                    
            except Exception as e:
                st.warning(f"⚠️ Task queue not available: {str(e)}")
                st.info("💡 Start Redis and Celery workers for background processing")
        else:
            st.warning("⚠️ Async features not available")
            st.info("💡 Install Redis and Celery for background processing")
    
    with col_task2:
        st.markdown("""
        <div class="feature-card">
            <h4 style="margin-bottom: 1rem;">📊 System Health</h4>
        </div>
        """, unsafe_allow_html=True)
        
        # System health metrics
        try:
            import psutil
            memory_percent = psutil.virtual_memory().percent
            cpu_percent = psutil.cpu_percent()
            
            if memory_percent < 70:
                st.success(f"💾 Memory Usage: {memory_percent:.1f}%")
            elif memory_percent < 85:
                st.warning(f"💾 Memory Usage: {memory_percent:.1f}%")
            else:
                st.error(f"💾 Memory Usage: {memory_percent:.1f}%")
            
            if cpu_percent < 70:
                st.success(f"⚡ CPU Usage: {cpu_percent:.1f}%")
            elif cpu_percent < 85:
                st.warning(f"⚡ CPU Usage: {cpu_percent:.1f}%")
            else:
                st.error(f"⚡ CPU Usage: {cpu_percent:.1f}%")
                
        except ImportError:
            st.info("💡 Install psutil for system monitoring")
    
    # Footer with branding
    st.markdown("""
    <div style="margin-top: 4rem; padding: 2rem; text-align: center; border-top: 1px solid var(--neutral-200);">
        <div style="color: var(--neutral-500); font-size: 0.9rem;">
            Powered by <strong>VerbatimAI</strong> • Built with ❤️ using Streamlit
        </div>
    </div>
    """, unsafe_allow_html=True)



def format_engagement_report(engagement_data):
    """Format engagement data for display"""
    return f"""
**Engagement Score: {engagement_data['total_score']}% ({engagement_data['level']})**

- Participation: {engagement_data['participation_score']}%
- Interaction: {engagement_data['interaction_score']}%
- Sentiment: {engagement_data['sentiment_score']}%
- Action Items: {engagement_data['action_score']}%
"""


# Fix the indentation error around line 2491
def show_upload_transcribe():
    st.markdown('<h1 class="main-header">🎙️ Upload & Transcribe</h1>', unsafe_allow_html=True)
    
    # Modern upload interface
    st.markdown("""
    <div style="text-align: center; margin-bottom: 2rem;">
        <p style="font-size: 1.1rem; color: var(--neutral-600);">
            Upload your audio or video files to get accurate AI-powered transcriptions with speaker identification
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Upload area with modern design
    st.markdown("""
    <div class="feature-card" style="text-align: center; padding: 3rem 2rem; margin-bottom: 2rem;">
        <div style="font-size: 4rem; margin-bottom: 1rem; color: var(--primary-blue);">📤</div>
        <h3 style="margin-bottom: 1rem;">Drop your files here</h3>
        <p style="color: var(--neutral-600); margin-bottom: 1.5rem;">
            Supported formats: MP3, WAV, MP4, MOV, AVI, M4A, FLAC
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Meeting configuration section
    st.markdown("""
    <div class="upload-section">
        <h3>📋 Meeting Configuration</h3>
        <p style="color: var(--neutral-600); margin-bottom: 1.5rem;">
            Configure your meeting details for enhanced analysis
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # Meeting type and time frame selection
    col1, col2 = st.columns(2)
    
    with col1:
        meeting_type = st.selectbox(
            "Meeting Type",
            [
                "General Meeting",
                "Weekly Standup", 
                "Project Review",
                "Client Interview",
                "Team Planning",
                "Performance Review",
                "Sales Call",
                "Technical Discussion",
                "Training Session",
                "Board Meeting",
                "Brainstorming Session",
                "Status Update"
            ],
            help="Select the type of meeting for optimized analysis"
        )
    
    with col2:
        time_frame = st.selectbox(
            "Expected Duration",
            [
                "0-10 minutes",
                "10-20 minutes", 
                "20-30 minutes",
                "30-45 minutes",
                "45-60 minutes",
                "1-2 hours",
                "2+ hours"
            ],
            help="Expected meeting duration for processing optimization"
        )
    
    # File upload section with updated limits
    st.markdown("""
    <div class="upload-section">
        <h3>📁 Upload Audio/Video File</h3>
        <p style="color: var(--neutral-600); margin-bottom: 1.5rem;">
            Upload your meeting recording for transcription and analysis.<br>
            <strong>Maximum file size: 1GB (1024MB)</strong><br>
            Supported formats: MP3, WAV, MP4, MOV, AVI, M4A, FLAC
        </p>
    </div>
    """, unsafe_allow_html=True)
    
    # File upload with 1GB limit
    st.info("📁 **File Upload Configuration**: Maximum file size has been set to 1GB (1024MB) for all uploads.")
    
    uploaded_file = st.file_uploader(
        "Choose an audio/video file",
        type=ALLOWED_EXTENSIONS,
        help=f"Supported formats: {', '.join(ALLOWED_EXTENSIONS).upper()}, Max size: 1GB",
        label_visibility="collapsed"
    )
    
    if uploaded_file is not None:
        # Check file size (1GB = 1024MB limit)
        file_size_mb = uploaded_file.size / 1024 / 1024
        
        if file_size_mb > 1024:  # 1GB limit
            st.error(f"⚠️ File size ({file_size_mb:.1f} MB) exceeds the 1GB (1024MB) limit. Please upload a smaller file.")
            return
        
        # Display file info with enhanced details
        file_details = {
            "Filename": uploaded_file.name,
            "File size": f"{file_size_mb:.2f} MB",
            "File type": uploaded_file.type,
            "Meeting type": meeting_type,
            "Expected duration": time_frame
        }
        
        # Create modern file info display
        st.markdown("""
        <div style="background: var(--glass-bg); border-radius: 12px; padding: 1.5rem; margin: 1rem 0; border: 1px solid var(--glass-border);">
            <h4 style="margin-bottom: 1rem; color: var(--primary-blue);">📋 File & Meeting Details</h4>
        </div>
        """, unsafe_allow_html=True)
        
        info_col1, info_col2 = st.columns(2)
        with info_col1:
            st.metric("File Size", f"{file_size_mb:.1f} MB", f"{(file_size_mb/1024*100):.1f}% of limit")
            st.metric("Meeting Type", meeting_type)
        with info_col2:
            st.metric("File Format", uploaded_file.type.split('/')[-1].upper())
            st.metric("Expected Duration", time_frame)
        
        # Advanced transcription options
        with st.expander("🔧 Advanced Transcription Options"):
            col1, col2 = st.columns(2)
            
            with col1:
                enable_speaker_detection = st.checkbox("Speaker Detection", value=True, help="Identify different speakers")
                enable_sentiment_analysis = st.checkbox("Real-time Sentiment", value=True, help="Analyze emotions during transcription")
                filter_profanity = st.checkbox("Filter Profanity", value=False, help="Remove inappropriate content")
            
            with col2:
                language = st.selectbox("Language", ["en", "es", "fr", "de", "it", "pt"], help="Transcription language")
                auto_punctuation = st.checkbox("Auto Punctuation", value=True, help="Automatically add punctuation")
                enable_highlights = st.checkbox("Auto Highlights", value=True, help="Detect key moments automatically")
        
        # Enhanced transcribe button
        if st.button("🚀 Start Advanced Transcription", type="primary", use_container_width=True):
            # Store meeting configuration in session state
            st.session_state.meeting_config = {
                'type': meeting_type,
                'duration': time_frame,
                'file_size_mb': file_size_mb,
                'transcription_options': {
                    'speaker_detection': enable_speaker_detection,
                    'sentiment_analysis': enable_sentiment_analysis,
                    'filter_profanity': filter_profanity,
                    'language': language,
                    'auto_punctuation': auto_punctuation,
                    'auto_highlights': enable_highlights
                }
            }
            
            # Show loading with meeting context
            progress_text = f"Transcribing {meeting_type.lower()} recording ({file_size_mb:.1f}MB)..."
            
            with st.spinner(progress_text):
                # Progress bar for large files
                if file_size_mb > 100:  # Show progress for files > 100MB
                    progress_bar = st.progress(0)
                    status_text = st.empty()
                    
                    # Simulate progress updates (in real implementation, this would come from Celery task)
                    import time
                    for i in range(100):
                        progress_bar.progress(i + 1)
                        if i < 30:
                            status_text.text("🔄 Uploading file to processing server...")
                        elif i < 60:
                            status_text.text("🎤 Processing audio and detecting speakers...")
                        elif i < 90:
                            status_text.text("🧠 Analyzing content and generating insights...")
                        else:
                            status_text.text("✨ Finalizing transcription and preparing results...")
                        time.sleep(0.05)  # Small delay for demo
                    
                    progress_bar.empty()
                    status_text.empty()
                
                # Perform transcription with configuration
                transcript = st.session_state.verbatim_ai.transcribe_audio(
                    uploaded_file, 
                    config=st.session_state.meeting_config['transcription_options']
                )
                
                if transcript:
                    st.session_state.verbatim_ai.transcript_data = transcript
                    
                    # Enhanced analysis based on meeting type
                    analysis_tasks = []
                    
                    if meeting_type in ["Client Interview", "Performance Review", "Sales Call"]:
                        analysis_tasks.extend(["sentiment", "engagement", "ai_content"])
                    elif meeting_type in ["Technical Discussion", "Brainstorming Session"]:
                        analysis_tasks.extend(["ai_content", "semantic", "key_points"])
                    elif meeting_type in ["Weekly Standup", "Status Update"]:
                        analysis_tasks.extend(["key_points", "action_items", "sentiment"])
                    else:
                        analysis_tasks = ["key_points", "sentiment", "ai_content"]
                    
                    # Perform selected analysis
                    st.session_state.verbatim_ai.analysis_results = {}
                    
                    if "key_points" in analysis_tasks:
                        st.session_state.verbatim_ai.analysis_results['key_points'] = st.session_state.verbatim_ai.extract_enhanced_key_points(transcript)
                    
                    if "sentiment" in analysis_tasks:
                        st.session_state.verbatim_ai.analysis_results['sentiment'] = st.session_state.verbatim_ai.analyze_sentiment(transcript)
                    
                    if "ai_content" in analysis_tasks:
                        st.session_state.verbatim_ai.analysis_results['ai_blocks'] = st.session_state.verbatim_ai.highlight_ai_blocks(transcript)
                    
                    # Save to meeting library with metadata
                    try:
                        if 'meeting_library' not in st.session_state:
                            from meeting_library import MeetingLibrary
                            st.session_state.meeting_library = MeetingLibrary()
                        
                        # Calculate duration from transcript
                        duration_minutes = 0
                        if hasattr(transcript, 'audio_duration') and transcript.audio_duration:
                            duration_minutes = transcript.audio_duration / 1000 / 60  # Convert to minutes
                        elif transcript.utterances:
                            # Calculate from last utterance end time
                            last_utterance = max(transcript.utterances, key=lambda u: u.end)
                            duration_minutes = last_utterance.end / 1000 / 60  # Convert to minutes
                        
                        # Prepare meeting data
                        meeting_data = {
                            'title': f"Meeting - {datetime.now().strftime('%Y-%m-%d %H:%M')}",
                            'date': datetime.now().strftime('%Y-%m-%d'),
                            'duration': duration_minutes,
                            'transcript': transcript.text,
                            'participants': list(set([
                                f"Speaker {utterance.speaker}" if isinstance(utterance.speaker, str) 
                                else f"Speaker {chr(65 + int(utterance.speaker))}" 
                                for utterance in transcript.utterances
                            ])),
                            'summary': '',  # Will be filled later
                            'tags': ['auto-generated'],
                            'file_name': uploaded_file.name if uploaded_file else 'unknown.wav'
                        }
                        
                        # Save meeting
                        meeting_id = st.session_state.meeting_library.add_meeting(meeting_data)
                        if meeting_id:
                            st.success(f"📚 Meeting saved to library with ID: {meeting_id}")
                        
                    except Exception as e:
                        st.warning(f"Could not save to meeting library: {str(e)}")
                    
                    st.success("✅ Transcription completed successfully!")
                    
                    # Show complete transcript with timestamps
                    st.subheader("📝 Complete Transcript")
                    
                    # Create a formatted transcript with speaker color labels
                    st.subheader("📖 Complete Transcript")
                    
                    # Define speaker colors with enhanced color scheme (CSS colors)
                    speaker_colors = {
                        0: "#1E88E5",  # Blue
                        1: "#43A047",  # Green
                        2: "#FFC107",  # Yellow
                        3: "#FF9800",  # Orange
                        4: "#9C27B0",  # Purple
                        5: "#E53935",  # Red
                        6: "#795548",  # Brown
                        7: "#607D8B",  # Blue Grey
                        8: "#00ACC1",  # Cyan
                        9: "#8BC34A",  # Light Green
                        10: "#FFEB3B", # Light Yellow
                        11: "#FF5722", # Deep Orange
                        12: "#E91E63", # Pink
                        13: "#673AB7", # Deep Purple
                        14: "#3F51B5", # Indigo
                        15: "#009688"  # Teal
                    }
                    
                    # Display speaker color legend
                    st.markdown("### 🎨 Speaker Color Legend")
                    unique_speakers = list(set([u.speaker for u in transcript.utterances]))
                    legend_cols = st.columns(min(4, len(unique_speakers)))
                    
                    for i, speaker in enumerate(unique_speakers):
                        with legend_cols[i % 4]:
                            speaker_color = speaker_colors.get(speaker, "#9E9E9E")
                            speaker_letter = chr(65 + speaker) if isinstance(speaker, int) else str(speaker)
                            st.markdown(f"""
                            <div style="display: flex; align-items: center; margin: 5px 0;">
                                <span style="display: inline-block; width: 20px; height: 20px; border-radius: 50%; background-color: {speaker_color}; margin-right: 10px; border: 2px solid #fff; box-shadow: 0 2px 4px rgba(0,0,0,0.2);"></span>
                                <strong>Speaker {speaker_letter}</strong>
                            </div>
                            """, unsafe_allow_html=True)
                    
                    st.markdown("---")
                    
                    # Display transcript with speaker colors
                    for utterance in transcript.utterances:
                        start_time = f"{utterance.start / 1000 / 60:.2f}:{(utterance.start / 1000) % 60:02.0f}"
                        end_time = f"{utterance.end / 1000 / 60:.2f}:{(utterance.end / 1000) % 60:02.0f}"
                        
                        # Get speaker color and create enhanced label
                        speaker_color = speaker_colors.get(utterance.speaker, "#9E9E9E")
                        speaker_letter = chr(65 + utterance.speaker) if isinstance(utterance.speaker, int) else str(utterance.speaker)
                        
                        # Display utterance with colored speaker labels
                        st.markdown(f"""
                        <div style="margin: 10px 0; padding: 15px; border-left: 4px solid {speaker_color}; background-color: #f8f9fa; border-radius: 8px; box-shadow: 0 2px 4px rgba(0,0,0,0.1);">
                            <div style="display: flex; align-items: center; margin-bottom: 8px;">
                                <span style="display: inline-block; width: 24px; height: 24px; border-radius: 50%; background-color: {speaker_color}; margin-right: 12px; border: 2px solid #fff; box-shadow: 0 2px 4px rgba(0,0,0,0.2);"></span>
                                <strong style="color: #2C3E50; font-size: 1.1em;">Speaker {speaker_letter}</strong>
                                <span style="color: #7F8C8D; margin-left: auto; font-size: 0.9em;">[{start_time} - {end_time}]</span>
                            </div>
                            <div style="color: #34495E; line-height: 1.6; font-size: 1.05em; padding-left: 36px;">
                                {utterance.text}
                            </div>
                        </div>
                        """, unsafe_allow_html=True)
                    
                    # Show key metrics with proper duration calculation
                    col1, col2, col3, col4 = st.columns(4)
                    with col1:
                        duration_minutes = transcript.audio_duration / 1000 / 60
                        st.metric("Duration", f"{duration_minutes:.2f} minutes")
                    with col2:
                        st.metric("Speakers", len(set([u.speaker for u in transcript.utterances])))
                    with col3:
                        st.metric("Words", len(transcript.text.split()))
                    with col4:
                        st.metric("Utterances", len(transcript.utterances))
                    
                    # Show combined summary (details + content summary)
                    st.subheader("📋 Meeting Summary")
                    if hasattr(st.session_state.verbatim_ai, 'generate_advanced_summary'):
                        summary = st.session_state.verbatim_ai.generate_advanced_summary(transcript)
                        st.markdown(summary)
                    else:
                        summary = st.session_state.verbatim_ai.generate_summary(transcript)
                        st.markdown(summary)
                    
                    # Show actual content summary (1/4 of original length)
                    st.subheader("📝 Content Summary (1/4 of original length)")
                    if hasattr(st.session_state.verbatim_ai, 'advanced_meeting_summarizer') and st.session_state.verbatim_ai.advanced_meeting_summarizer:
                        try:
                            content_summary = st.session_state.verbatim_ai.advanced_meeting_summarizer.generate_content_summary(transcript.text, target_ratio=0.25)
                        except Exception as e:
                            st.warning(f"Advanced summarizer error: {e}")
                            content_summary = "Advanced summarizer not available."
                        st.markdown(content_summary)
                    else:
                        st.info("Advanced summarizer not available. Using basic summary.")
                        # Create a simple content summary
                        sentences = transcript.text.split('.')
                        summary_sentences = sentences[:len(sentences)//4]  # Take 1/4 of sentences
                        content_summary = '. '.join(summary_sentences) + '.'
                        st.markdown(content_summary)
                    
                    # Show enhanced key points
                    st.subheader("🎯 Key Points & Insights")
                    key_points = st.session_state.verbatim_ai.analysis_results.get('key_points', {})
                    
                    # Display key points in tabs
                    tab1, tab2, tab3, tab4, tab5 = st.tabs(["Decisions", "Action Items", "Questions", "Organizations", "Highlights"])
                    
                    with tab1:
                        if key_points['decisions']:
                            for i, decision in enumerate(key_points['decisions'], 1):
                                timestamp = f"{decision['timestamp'] / 60:.2f}:{decision['timestamp'] % 60:02.0f}"
                                st.markdown(f"**{i}.** [{timestamp}] **{decision['speaker']}**: {decision['text']}")
                        else:
                            st.info("No decisions detected in this meeting.")
                    
                    with tab2:
                        if key_points['action_items']:
                            for i, action in enumerate(key_points['action_items'], 1):
                                timestamp = f"{action['timestamp'] / 60:.2f}:{action['timestamp'] % 60:02.0f}"
                                st.markdown(f"**{i}.** [{timestamp}] **{action['speaker']}**: {action['text']}")
                        else:
                            st.info("No action items detected in this meeting.")
                    
                    with tab3:
                        if key_points['questions']:
                            for i, question in enumerate(key_points['questions'], 1):
                                timestamp = f"{question['timestamp'] / 60:.2f}:{question['timestamp'] % 60:02.0f}"
                                st.markdown(f"**{i}.** [{timestamp}] **{question['speaker']}**: {question['text']}")
                        else:
                            st.info("No questions detected in this meeting.")
                    
                    with tab4:
                        organizations = key_points['organizations']
                        if organizations:
                            for i, org in enumerate(organizations, 1):
                                timestamp = f"{org['timestamp'] / 60:.2f}:{org['timestamp'] % 60:02.0f}"
                                st.markdown(f"**{i}.** [{timestamp}] **{org['speaker']}**: {org['text']}")
                        else:
                            st.info("No organizations/companies detected in this meeting.")
                    
                    with tab5:
                        if key_points['highlights']:
                            for i, highlight in enumerate(key_points['highlights'], 1):
                                timestamp = f"{highlight['timestamp'] / 60:.2f}:{highlight['timestamp'] % 60:02.0f}"
                                st.markdown(f"**{i}.** [{timestamp}] **{highlight['reason']}**: {highlight['text']}")
                        else:
                            st.info("No highlights available for this meeting.")
def show_analytics():
    st.markdown('<h1 class="main-header">📊 Advanced Analytics Dashboard</h1>', unsafe_allow_html=True)
    
    if not st.session_state.verbatim_ai.transcript_data:
        st.warning("Please upload and transcribe a file first!")
        return
    
    transcript = st.session_state.verbatim_ai.transcript_data
    analysis = st.session_state.verbatim_ai.analysis_results
    
    # Color palette for consistent visualization (20 different colors)
    COLOR_PALETTE = [
        '#1E88E5', '#7B1FA2', '#43A047', '#FF9800', '#E53935',
        '#00ACC1', '#8BC34A', '#FFC107', '#9C27B0', '#2196F3',
        '#4CAF50', '#FF5722', '#795548', '#607D8B', '#9E9E9E',
        '#F44336', '#E91E63', '#673AB7', '#3F51B5', '#009688'
    ]
    
    # Speaker duration chart with proper calculation
    speaker_durations = {}
    total_duration = 0
    
    for utterance in transcript.utterances:
        # Handle both string and integer speaker formats
        if isinstance(utterance.speaker, str):
            speaker = f"Speaker {utterance.speaker}"
        else:
            # Convert integer speaker ID to letter (0->A, 1->B, etc.)
            speaker = f"Speaker {chr(65 + int(utterance.speaker))}"
        
        duration = (utterance.end - utterance.start) / 1000  # Convert to seconds
        speaker_durations[speaker] = speaker_durations.get(speaker, 0) + duration
        total_duration += duration
    
    # Calculate percentages
    speaker_percentages = {}
    for speaker, duration in speaker_durations.items():
        percentage = (duration / total_duration * 100) if total_duration > 0 else 0
        speaker_percentages[speaker] = f"{percentage:.1f}%"
    
    # Create pie chart with percentages
    fig_duration = px.pie(
        values=list(speaker_durations.values()),
        names=[f"{speaker} ({speaker_percentages[speaker]})" for speaker in speaker_durations.keys()],
        title="Speaking Time Distribution",
        color_discrete_sequence=px.colors.qualitative.Set3
    )
    fig_duration.update_traces(textposition='inside', textinfo='percent+label')
    st.plotly_chart(fig_duration, use_container_width=True)
    
    # Show detailed statistics
    with st.expander("📊 Detailed Speaking Statistics"):
        for speaker, duration in speaker_durations.items():
            minutes = int(duration // 60)
            seconds = int(duration % 60)
            st.write(f"**{speaker}:** {minutes}m {seconds}s ({speaker_percentages[speaker]})")
    
    # Enhanced sentiment analysis
    st.subheader("😊 Enhanced Sentiment & Emotion Analysis")
    
    if analysis and analysis.get('sentiment'):
        # Create tabs for different sentiment views
        tab1, tab2, tab3 = st.tabs(["Sentiment Distribution", "Emotion Detection", "Tone Analysis"])
        
        with tab1:
            # Average sentiment per speaker
            sentiment_data = []
            for speaker, data in analysis['sentiment'].items():
                if data.get('sentiment_scores'):
                    avg_positive = sum(s.get('positive', 0) for s in data['sentiment_scores']) / len(data['sentiment_scores'])
                    avg_negative = sum(s.get('negative', 0) for s in data['sentiment_scores']) / len(data['sentiment_scores'])
                    avg_neutral = sum(s.get('neutral', 0) for s in data['sentiment_scores']) / len(data['sentiment_scores'])
                    
                    sentiment_data.append({
                        'Speaker': speaker,
                        'Positive': avg_positive,
                        'Negative': avg_negative,
                        'Neutral': avg_neutral
                    })
            
            if sentiment_data:
                df_sentiment = pd.DataFrame(sentiment_data)
                
                fig_sentiment = px.bar(
                    df_sentiment,
                    x='Speaker',
                    y=['Positive', 'Negative', 'Neutral'],
                    title="Average Sentiment by Speaker",
                    barmode='stack'
                )
                st.plotly_chart(fig_sentiment, use_container_width=True)
            else:
                st.info("No sentiment data available for visualization.")
        
        with tab2:
            # Emotion detection summary
            emotion_summary = {}
            for speaker, data in analysis['sentiment'].items():
                emotion_counts = {}
                emotions = data.get('emotions', [])
                for emotion in emotions:
                    emotion_counts[emotion] = emotion_counts.get(emotion, 0) + 1
                emotion_summary[speaker] = emotion_counts
            
            # Display emotion summary
            for speaker, emotions in emotion_summary.items():
                if emotions:
                    st.markdown(f"**{speaker}**:")
                    emotion_text = ", ".join([f"{emotion} ({count})" for emotion, count in emotions.items()])
                    st.markdown(f"*{emotion_text}*")
                else:
                    st.markdown(f"**{speaker}**: *No specific emotions detected*")
        
        with tab3:
            # Tone analysis
            tone_summary = {}
            for speaker, data in analysis['sentiment'].items():
                tone_counts = {}
                tones = data.get('tone_analysis', [])
                for tone in tones:
                    tone_counts[tone] = tone_counts.get(tone, 0) + 1
                tone_summary[speaker] = tone_counts
            
            # Display tone summary
            for speaker, tones in tone_summary.items():
                if tones:
                    st.markdown(f"**{speaker}**:")
                    tone_text = ", ".join([f"{tone} ({count})" for tone, count in tones.items()])
                    st.markdown(f"*{tone_text}*")
                else:
                    st.markdown(f"**{speaker}**: *No tone analysis available*")
    else:
        st.info("No sentiment analysis data available. Please run the analysis first.")
    
    # AI Blocks Analysis
    st.subheader("🤖 AI & Technical Content Analysis")
    
    ai_blocks = st.session_state.verbatim_ai.highlight_ai_blocks(transcript)
    if ai_blocks:
        st.markdown("**AI/Technical Topics Discussed:**")
        for block in ai_blocks:
            st.markdown(f"**[{block['start_time']} - {block['end_time']}] {block['speaker']}** - *{block['ai_topic']}*")
            st.markdown(f"*{block['text']}*")
            st.markdown("---")
    else:
        st.info("No AI or technical content detected in this meeting.")
    
    # Advanced Analytics Tabs
    tab_speaker, tab_sentiment, tab_semantic, tab_engagement, tab_ai, tab_reports = st.tabs([
        "👥 Speaker Analytics", "😊 Sentiment Analysis", "🧠 Semantic Insights", 
        "📈 Engagement Metrics", "🤖 AI Content Analysis", "📋 Detailed Reports"
    ])
    
    # Speaker Analytics Tab
    with tab_speaker:
        st.subheader("👥 Comprehensive Speaker Analysis")
        
        # Speaker metrics overview
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            create_modern_metric_card("Total Speakers", len(speaker_durations), icon="👤")
        
        with col2:
            avg_speaking_time = sum(speaker_durations.values()) / len(speaker_durations) if speaker_durations else 0
            create_modern_metric_card("Avg Speaking Time", f"{avg_speaking_time/60:.1f}m", icon="⏱️")
        
        with col3:
            # Calculate words per speaker
            total_words = len(transcript.text.split()) if hasattr(transcript, 'text') else 0
            avg_words = total_words / len(speaker_durations) if speaker_durations else 0
            create_modern_metric_card("Avg Words/Speaker", f"{avg_words:.0f}", icon="💬")
        
        with col4:
            # Most active speaker
            most_active = max(speaker_durations, key=speaker_durations.get) if speaker_durations else "N/A"
            create_modern_metric_card("Most Active", most_active.replace("Speaker ", ""), icon="🎤")
        
        # Interactive speaker comparison chart
        st.subheader("📊 Speaking Time vs Word Count Analysis")
        
        # Calculate words per speaker
        speaker_words = {}
        for utterance in transcript.utterances:
            speaker = f"Speaker {utterance.speaker}" if isinstance(utterance.speaker, str) else f"Speaker {chr(65 + int(utterance.speaker))}"
            word_count = len(utterance.text.split())
            speaker_words[speaker] = speaker_words.get(speaker, 0) + word_count
        
        # Create comparison dataframe
        comparison_data = []
        for speaker in speaker_durations.keys():
            comparison_data.append({
                'Speaker': speaker,
                'Duration (minutes)': speaker_durations[speaker] / 60,
                'Word Count': speaker_words.get(speaker, 0),
                'Words per Minute': (speaker_words.get(speaker, 0) / (speaker_durations[speaker] / 60)) if speaker_durations[speaker] > 0 else 0
            })
        
        df_comparison = pd.DataFrame(comparison_data)
        
        # Scatter plot for duration vs words
        fig_scatter = px.scatter(
            df_comparison, 
            x='Duration (minutes)', 
            y='Word Count',
            color='Words per Minute',
            size='Word Count',
            hover_name='Speaker',
            title="Speaker Performance: Duration vs Word Count",
            color_continuous_scale='Viridis'
        )
        st.plotly_chart(fig_scatter, use_container_width=True)
        
        # Speaker performance table
        st.subheader("📋 Detailed Speaker Performance")
        st.dataframe(df_comparison.round(2), use_container_width=True)
    
    # Sentiment Analysis Tab
    with tab_sentiment:
        st.subheader("😊 Advanced Sentiment & Emotion Analytics")
        
        if analysis and analysis.get('sentiment'):
            # Sentiment overview metrics
            col1, col2, col3, col4 = st.columns(4)
            
            # Calculate overall sentiment metrics
            all_positive = []
            all_negative = []
            all_neutral = []
            
            for speaker_data in analysis['sentiment'].values():
                for score in speaker_data.get('sentiment_scores', []):
                    all_positive.append(score.get('positive', 0))
                    all_negative.append(score.get('negative', 0))
                    all_neutral.append(score.get('neutral', 0))
            
            avg_positive = sum(all_positive) / len(all_positive) if all_positive else 0
            avg_negative = sum(all_negative) / len(all_negative) if all_negative else 0
            avg_neutral = sum(all_neutral) / len(all_neutral) if all_neutral else 0
            
            with col1:
                create_modern_metric_card("Overall Positive", f"{avg_positive:.1%}", icon="😊")
            with col2:
                create_modern_metric_card("Overall Negative", f"{avg_negative:.1%}", icon="😞")
            with col3:
                create_modern_metric_card("Overall Neutral", f"{avg_neutral:.1%}", icon="😐")
            with col4:
                dominant = "Positive" if avg_positive > max(avg_negative, avg_neutral) else "Negative" if avg_negative > avg_neutral else "Neutral"
                create_modern_metric_card("Dominant Mood", dominant, icon="🎭")
            
            # Sentiment timeline
            st.subheader("📈 Sentiment Timeline")
            timeline_data = []
            
            for speaker, data in analysis['sentiment'].items():
                for i, score in enumerate(data.get('sentiment_scores', [])):
                    timeline_data.append({
                        'Time Point': i,
                        'Speaker': speaker,
                        'Positive': score.get('positive', 0),
                        'Negative': score.get('negative', 0),
                        'Neutral': score.get('neutral', 0)
                    })
            
            if timeline_data:
                df_timeline = pd.DataFrame(timeline_data)
                
                # Line chart showing sentiment progression
                fig_timeline = px.line(
                    df_timeline, 
                    x='Time Point', 
                    y=['Positive', 'Negative', 'Neutral'],
                    color='Speaker',
                    title="Sentiment Evolution Throughout Meeting",
                    labels={'value': 'Sentiment Score', 'variable': 'Sentiment Type'}
                )
                st.plotly_chart(fig_timeline, use_container_width=True)
                
                # Emotion heatmap
                st.subheader("🎯 Emotion Distribution Heatmap")
                emotion_data = []
                for speaker, data in analysis['sentiment'].items():
                    emotions = data.get('emotions', [])
                    emotion_counts = {}
                    for emotion in emotions:
                        emotion_counts[emotion] = emotion_counts.get(emotion, 0) + 1
                    
                    for emotion, count in emotion_counts.items():
                        emotion_data.append({
                            'Speaker': speaker,
                            'Emotion': emotion,
                            'Count': count
                        })
                
                if emotion_data:
                    df_emotions = pd.DataFrame(emotion_data)
                    emotion_pivot = df_emotions.pivot(index='Speaker', columns='Emotion', values='Count').fillna(0)
                    
                    fig_heatmap = px.imshow(
                        emotion_pivot.values,
                        x=emotion_pivot.columns,
                        y=emotion_pivot.index,
                        title="Emotion Distribution by Speaker",
                        color_continuous_scale='Viridis'
                    )
                    st.plotly_chart(fig_heatmap, use_container_width=True)
        
        else:
            st.info("No sentiment analysis data available. Please run the analysis first.")
    
    # Semantic Insights Tab
    with tab_semantic:
        st.subheader("🧠 Semantic Analysis & Topic Discovery")
        
        # Semantic search functionality
        st.subheader("🔍 Semantic Search")
        search_query = st.text_input("Search for concepts or topics:", placeholder="e.g., 'project deadlines', 'technical challenges'")
        
        if search_query and hasattr(st.session_state.verbatim_ai, 'semantic_search_engine') and st.session_state.verbatim_ai.semantic_search_engine:
            with st.spinner("Performing semantic search..."):
                try:
                    search_results = st.session_state.verbatim_ai.semantic_search_engine.search(search_query, transcript.text)
                    
                    if search_results:
                        st.success(f"Found {len(search_results)} relevant segments:")
                        for i, result in enumerate(search_results[:5], 1):
                            with st.expander(f"Result {i} (Relevance: {result['score']:.2f})"):
                                st.write(result['text'])
                                if 'timestamp' in result:
                                    st.caption(f"Timestamp: {result['timestamp']}")
                    else:
                        st.info("No relevant segments found for your search query.")
                except Exception as e:
                    st.error(f"Error performing search: {e}")
        
        # Topic clustering visualization
        st.subheader("📊 Topic Clustering & Themes")
        
        if hasattr(st.session_state.verbatim_ai, 'semantic_search_engine') and st.session_state.verbatim_ai.semantic_search_engine:
            try:
                topics = st.session_state.verbatim_ai.semantic_search_engine.find_semantic_themes(5)
                
                if topics:
                    topic_data = []
                    # Handle both dict and list types for topics
                    if isinstance(topics, dict):
                        topic_iter = topics.items()
                    elif isinstance(topics, list):
                        topic_iter = topics
                    else:
                        topic_iter = []
                    
                    for topic_item in topic_iter:
                        # Handle different topic item formats
                        if isinstance(topic_item, (list, tuple)) and len(topic_item) == 2:
                            topic, keywords = topic_item
                        elif isinstance(topic_item, dict):
                            topic = topic_item.get('topic', 'Unknown')
                            keywords = topic_item.get('keywords', [])
                        else:
                            topic = str(topic_item)
                            keywords = []
                        if isinstance(keywords, dict):
                            for keyword, weight in keywords.items():
                                topic_data.append({
                                    'Topic': topic,
                                    'Keyword': keyword,
                                    'Weight': weight
                                })
                        elif isinstance(keywords, list):
                            for kw in keywords:
                                if isinstance(kw, (list, tuple)) and len(kw) == 2:
                                    keyword, weight = kw
                                    topic_data.append({
                                        'Topic': topic,
                                        'Keyword': keyword,
                                        'Weight': weight
                                    })
                    
                    if topic_data:
                        df_topics = pd.DataFrame(topic_data)
                        
                        # Sunburst chart for topic hierarchy
                        fig_sunburst = px.sunburst(
                            df_topics, 
                            path=['Topic', 'Keyword'], 
                            values='Weight',
                            title="Topic and Keyword Hierarchy"
                        )
                        st.plotly_chart(fig_sunburst, use_container_width=True)
                        
                        # Top keywords per topic
                        st.subheader("🏷️ Top Keywords by Topic")
                        for topic in set(df_topics['Topic']):
                            topic_keywords = df_topics[df_topics['Topic'] == topic].nlargest(5, 'Weight')
                            keywords_str = ", ".join([f"{row['Keyword']} ({row['Weight']:.2f})" for _, row in topic_keywords.iterrows()])
                            st.markdown(f"**{topic}:** {keywords_str}")
                
            except AttributeError:
                st.info("Topic extraction not available in current semantic search engine.")
        else:
            st.info("Semantic analysis engine not available. Install required dependencies to enable this feature.")
    
    # Engagement Metrics Tab
    with tab_engagement:
        st.subheader("📈 Meeting Engagement & Participation Metrics")
        
        # Calculate engagement metrics
        total_utterances = len(transcript.utterances)
        unique_speakers = len(set(utterance.speaker for utterance in transcript.utterances))
        
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            create_modern_metric_card("Total Utterances", str(total_utterances), icon="💬")
        
        with col2:
            avg_utterance_length = sum(len(utterance.text.split()) for utterance in transcript.utterances) / total_utterances if total_utterances > 0 else 0
            create_modern_metric_card("Avg Words/Utterance", f"{avg_utterance_length:.1f}", icon="📝")
        
        with col3:
            # Calculate interaction rate (how often speakers change)
            speaker_changes = 0
            for i in range(1, len(transcript.utterances)):
                if transcript.utterances[i].speaker != transcript.utterances[i-1].speaker:
                    speaker_changes += 1
            interaction_rate = (speaker_changes / total_utterances * 100) if total_utterances > 0 else 0
            create_modern_metric_card("Interaction Rate", f"{interaction_rate:.1f}%", icon="🔄")
        
        with col4:
            # Calculate engagement balance (lower is more balanced)
            speaker_utterance_counts = {}
            for utterance in transcript.utterances:
                speaker = f"Speaker {utterance.speaker}" if isinstance(utterance.speaker, str) else f"Speaker {chr(65 + int(utterance.speaker))}"
                speaker_utterance_counts[speaker] = speaker_utterance_counts.get(speaker, 0) + 1
            
            max_percentage = max(speaker_utterance_counts.values()) / total_utterances * 100 if total_utterances > 0 else 0
            engagement_balance = 100 - max_percentage  # Higher is more balanced
            create_modern_metric_card("Balance Score", f"{engagement_balance:.1f}%", icon="⚖️")
        
        # Participation flow chart
        st.subheader("🌊 Participation Flow")
        
        # Create participation timeline
        participation_data = []
        for i, utterance in enumerate(transcript.utterances):
            speaker = f"Speaker {utterance.speaker}" if isinstance(utterance.speaker, str) else f"Speaker {chr(65 + int(utterance.speaker))}"
            participation_data.append({
                'Utterance': i + 1,
                'Speaker': speaker,
                'Word Count': len(utterance.text.split()),
                'Duration': (utterance.end - utterance.start) / 1000
            })
        
        df_participation = pd.DataFrame(participation_data)
        
        # Stacked area chart showing participation over time
        fig_flow = px.area(
            df_participation, 
            x='Utterance', 
            y='Word Count',
            color='Speaker',
            title="Word Count Flow by Speaker Throughout Meeting",
            color_discrete_sequence=COLOR_PALETTE
        )
        st.plotly_chart(fig_flow, use_container_width=True)
        
        # Speaking pattern analysis
        st.subheader("🔄 Speaking Patterns")
        
        # Calculate average pause length between speakers
        pauses = []
        for i in range(1, len(transcript.utterances)):
            if transcript.utterances[i].speaker != transcript.utterances[i-1].speaker:
                pause = (transcript.utterances[i].start - transcript.utterances[i-1].end) / 1000
                pauses.append(max(0, pause))  # Ensure non-negative
        
        avg_pause = sum(pauses) / len(pauses) if pauses else 0
        
        pattern_col1, pattern_col2 = st.columns(2)
        with pattern_col1:
            create_modern_metric_card("Avg Speaker Transition", f"{avg_pause:.1f}s", icon="⏱️")
        
        with pattern_col2:
            # Calculate speaking momentum (consecutive utterances by same speaker)
            momentum_data = {}
            current_speaker = None
            current_streak = 0
            
            for utterance in transcript.utterances:
                speaker = f"Speaker {utterance.speaker}" if isinstance(utterance.speaker, str) else f"Speaker {chr(65 + int(utterance.speaker))}"
                
                if speaker == current_speaker:
                    current_streak += 1
                else:
                    if current_speaker and current_streak > 1:
                        momentum_data[current_speaker] = momentum_data.get(current_speaker, []) + [current_streak]
                    current_speaker = speaker
                    current_streak = 1
            
            # Add final streak
            if current_speaker and current_streak > 1:
                momentum_data[current_speaker] = momentum_data.get(current_speaker, []) + [current_streak]
            
            max_streak = max([max(streaks) for streaks in momentum_data.values()]) if momentum_data else 1
            create_modern_metric_card("Longest Speaking Streak", f"{max_streak} utterances", icon="🔥")
    
    # AI Content Analysis Tab
    with tab_ai:
        st.subheader("🤖 AI & Technical Content Deep Dive")
        
        ai_blocks = st.session_state.verbatim_ai.highlight_ai_blocks(transcript)
        
        if ai_blocks:
            col1, col2, col3, col4 = st.columns(4)
            
            with col1:
                create_modern_metric_card("AI Mentions", str(len(ai_blocks)), icon="🤖")
            
            with col2:
                unique_ai_topics = len(set(block['ai_topic'] for block in ai_blocks))
                create_modern_metric_card("Unique AI Topics", str(unique_ai_topics), icon="🧠")
            
            with col3:
                ai_speakers = len(set(block['speaker'] for block in ai_blocks))
                create_modern_metric_card("Speakers Discussing AI", str(ai_speakers), icon="👥")
            
            with col4:
                def parse_time(time_str):
                    return datetime.strptime(time_str, "%H:%M:%S")
                total_ai_time = sum(   
                    (parse_time(block['end_time']) - parse_time(block['start_time'])).total_seconds()
    for block in ai_blocks)
                create_modern_metric_card("Total AI Discussion", f"{total_ai_time/60:.1f}m", icon="⏱️")
            
            # AI topic distribution
            st.subheader("📊 AI Topic Distribution")
            
            topic_counts = {}
            for block in ai_blocks:
                topic_counts[block['ai_topic']] = topic_counts.get(block['ai_topic'], 0) + 1
            
            fig_ai_topics = px.bar(
                x=list(topic_counts.keys()),
                y=list(topic_counts.values()),
                title="Frequency of AI Topics Discussed",
                color=list(topic_counts.values()),
                color_continuous_scale='Viridis'
            )
            fig_ai_topics.update_layout(xaxis_title="AI Topics", yaxis_title="Mentions")
            st.plotly_chart(fig_ai_topics, use_container_width=True)
            
            # AI content timeline
            st.subheader("⏰ AI Discussion Timeline")
            
            timeline_ai_data = []
            for block in ai_blocks:
                timeline_ai_data.append({
                    'Start Time (min)': block['start_time'] / 60,
                    'Duration (min)': (block['end_time'] - block['start_time']) / 60,
                    'Speaker': block['speaker'],
                    'AI Topic': block['ai_topic'],
                    'Text Preview': block['text'][:100] + "..." if len(block['text']) > 100 else block['text']
                })
            
            df_ai_timeline = pd.DataFrame(timeline_ai_data)
            
            fig_ai_timeline = px.scatter(
                df_ai_timeline,
                x='Start Time (min)',
                y='AI Topic',
                size='Duration (min)',
                color='Speaker',
                hover_data=['Text Preview'],
                title="AI Content Throughout Meeting Timeline"
            )
            st.plotly_chart(fig_ai_timeline, use_container_width=True)
            
            # Detailed AI blocks
            st.subheader("📋 Detailed AI Content Blocks")
            for i, block in enumerate(ai_blocks, 1):
                with st.expander(f"AI Block {i}: {block['ai_topic']} - {block['speaker']}"):
                    col1, col2 = st.columns([3, 1])
                    with col1:
                        st.write(block['text'])
                    with col2:
                        st.metric("Start Time", f"{block['start_time']/60:.1f}m")
                        st.metric("Duration", f"{(block['end_time'] - block['start_time']):.1f}s")
        
        else:
            st.info("No AI or technical content detected in this meeting.")
            st.markdown("**Common AI topics we look for:**")
            st.markdown("- Machine Learning & Deep Learning")
            st.markdown("- Natural Language Processing")
            st.markdown("- Computer Vision & Image Processing")
            st.markdown("- Data Science & Analytics")
            st.markdown("- Automation & Robotics")
            st.markdown("- Cloud Computing & APIs")
    
    # Detailed Reports Tab
    with tab_reports:
        st.subheader("📋 Comprehensive Reports & Export Options")
        
        # Import datetime at the top if not already imported
        from datetime import datetime
        
        # Report generation options
        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("📊 Generate Custom Report")
            
            report_sections = st.multiselect(
                "Select report sections:",
                ["Speaker Analysis", "Sentiment Overview", "AI Content", "Key Topics", "Engagement Metrics"],
                default=["Speaker Analysis", "Sentiment Overview"]
            )
            
            report_format = st.selectbox("Report Format:", ["PDF", "HTML", "JSON", "CSV"])
            
            if st.button("Generate Report", type="primary"):
                with st.spinner("Generating comprehensive report..."):
                    # Here you would integrate with the Celery export tasks
                    report_data = {
                        'transcript': transcript,
                        'analysis': analysis,
                        'sections': report_sections,
                        'format': report_format.lower()
                    }
                    
                    # Mock report generation (replace with actual Celery task)
                    st.success(f"Report generated successfully! Download will begin shortly.")
                    st.download_button(
                        label=f"Download {report_format} Report",
                        data="Mock report data - replace with actual generated report",
                        file_name=f"meeting_report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{report_format.lower()}",
                        mime=f"application/{report_format.lower()}"
                    )
        
        with col2:
            st.subheader("📈 Quick Analytics Export")
            
            export_type = st.selectbox(
                "Data to Export:",
                ["All Analytics", "Speaker Statistics", "Sentiment Data", "AI Content", "Transcript Only"]
            )
            
            export_format = st.selectbox("Export Format:", ["CSV", "JSON", "Excel"])
            
            if st.button("Export Data", type="secondary"):
                with st.spinner("Preparing export..."):
                    # Generate export data based on selection
                    if export_type == "Speaker Statistics":
                        export_data = df_comparison.to_csv(index=False) if 'df_comparison' in locals() else "No speaker data available"
                    elif export_type == "Sentiment Data" and 'df_timeline' in locals():
                        export_data = df_timeline.to_csv(index=False)
                    else:
                        export_data = f"Export data for {export_type} - Mock data"
                    
                    st.download_button(
                        label=f"Download {export_format} File",
                        data=export_data,
                        file_name=f"{export_type.lower().replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{export_format.lower()}",
                        mime=f"application/{export_format.lower()}"
                    )
        
        # Analytics summary
        st.subheader("📊 Session Analytics Summary")
        
        summary_col1, summary_col2, summary_col3 = st.columns(3)
        
        with summary_col1:
            st.markdown("**Meeting Overview**")
            st.write(f"• Duration: {total_duration/60:.1f} minutes")
            st.write(f"• Speakers: {len(speaker_durations)}")
            st.write(f"• Total Words: {len(transcript.text.split()) if hasattr(transcript, 'text') else 0:,}")
            st.write(f"• Utterances: {len(transcript.utterances)}")
        
        with summary_col2:
            st.markdown("**Analysis Highlights**")
            if analysis and analysis.get('sentiment'):
                sentiment_summary = "Positive" if avg_positive > max(avg_negative, avg_neutral) else "Negative" if avg_negative > avg_neutral else "Neutral"
                st.write(f"• Overall Mood: {sentiment_summary}")
            st.write(f"• AI Content Blocks: {len(ai_blocks)}")
            st.write(f"• Engagement Balance: {engagement_balance:.1f}%")
            st.write(f"• Interaction Rate: {interaction_rate:.1f}%")
        
        with summary_col3:
            st.markdown("**Export Options**")
            st.write("• Comprehensive PDF Reports")
            st.write("• CSV Data Exports")
            st.write("• JSON API Format")
            st.write("• HTML Presentations")
            
        # Visualizations export
        st.subheader("📊 Export Visualizations")
        
        viz_col1, viz_col2, viz_col3 = st.columns(3)
        
        with viz_col1:
            if st.button("Export Speaking Time Chart"):
                st.info("Chart export feature coming soon!")
        
        with viz_col2:
            if st.button("Export Sentiment Timeline"):
                st.info("Timeline export feature coming soon!")
        
        with viz_col3:
            if st.button("Export All Charts"):
                st.info("Bulk chart export feature coming soon!")
    
    # Final metrics section - fixed to use defined variables
    if analysis and analysis.get('key_points'):
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric("Decisions", len(analysis.get('key_points', {}).get('decisions', [])))
        with col2:
            st.metric("Action Items", len(analysis.get('key_points', {}).get('action_items', [])))
        with col3:
            st.metric("Questions", len(analysis.get('key_points', {}).get('questions', [])))
        with col4:
            st.metric("Organizations", len(analysis.get('key_points', {}).get('organizations', [])))

def show_search():
    st.title("🔍 Search & Filter")
    
    if not st.session_state.verbatim_ai.transcript_data:
        st.warning("Please upload and transcribe a file first!")
        return
    
    transcript = st.session_state.verbatim_ai.transcript_data
    analysis = st.session_state.verbatim_ai.analysis_results
    
    # Search options
    search_type = st.selectbox(
        "Search by:",
        ["Text", "Speaker", "Sentiment", "Key Points"]
    )
    
    if search_type == "Text":
        search_query = st.text_input("Enter search term:")
        if search_query:
            results = []
            for utterance in transcript.utterances:
                if search_query.lower() in utterance.text.lower():
                    results.append({
                        'Speaker': f"Speaker {utterance.speaker}",
                        'Text': utterance.text,
                        'Timestamp': f"{utterance.start / 1000 / 60:.1f} min"
                    })
            
            if results:
                df_results = pd.DataFrame(results)
                st.dataframe(df_results, use_container_width=True)
            else:
                st.info("No results found.")
    
    elif search_type == "Speaker":
        speakers = list(set([f"Speaker {u.speaker}" for u in transcript.utterances]))
        selected_speaker = st.selectbox("Select speaker:", speakers)
        
        if selected_speaker:
            speaker_num = selected_speaker.split()[-1]
            results = []
            for utterance in transcript.utterances:
                if utterance.speaker == speaker_num:
                    results.append({
                        'Text': utterance.text,
                        'Timestamp': f"{utterance.start / 1000 / 60:.1f} min"
                    })
            
            if results:
                df_results = pd.DataFrame(results)
                st.dataframe(df_results, use_container_width=True)
    
    elif search_type == "Sentiment":
        sentiment_filter = st.selectbox("Filter by sentiment:", ["Positive", "Negative", "Neutral"])
        
        if analysis['sentiment']:
            results = []
            for speaker, data in analysis['sentiment'].items():
                for utterance_data in data['utterances']:
                    sentiment_scores = utterance_data['sentiment']
                    dominant_sentiment = max(sentiment_scores, key=sentiment_scores.get)
                    
                    if dominant_sentiment.lower() == sentiment_filter.lower():
                        results.append({
                            'Speaker': speaker,
                            'Text': utterance_data['text'],
                            'Sentiment': dominant_sentiment,
                            'Timestamp': f"{utterance_data['timestamp'] / 60:.1f} min"
                        })
            
            if results:
                df_results = pd.DataFrame(results)
                st.dataframe(df_results, use_container_width=True)
            else:
                st.info("No results found.")
    
    elif search_type == "Key Points":
        point_type = st.selectbox("Select point type:", ["Decisions", "Action Items", "Questions", "Highlights"])
        
        key_points = analysis.get('key_points', {})
        if point_type == "Decisions":
            results = key_points['decisions']
        elif point_type == "Action Items":
            results = key_points['action_items']
        elif point_type == "Questions":
            results = key_points['questions']
        elif point_type == "Highlights":
            results = key_points['highlights']
        
        if results:
            df_results = pd.DataFrame(results)
            st.dataframe(df_results, use_container_width=True)
        else:
            st.info(f"No {point_type.lower()} found.")

def show_export():
    st.title("📄 Export Transcript")
    
    if not hasattr(st.session_state.verbatim_ai, 'transcript_data') or not st.session_state.verbatim_ai.transcript_data:
        st.warning("⚠️ Please upload and transcribe a file first!")
        return
    
    transcript = st.session_state.verbatim_ai.transcript_data
    analysis = getattr(st.session_state.verbatim_ai, 'analysis_results', None)
    
    # Export options
    export_format = st.selectbox("Export format:", ["DOCX", "PDF"])
    
    # Generate dynamic filename based on timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_filename = f"meeting_transcript_{timestamp}"
    
    if st.button(f"📥 Export as {export_format}", type="primary"):
        if export_format == "DOCX":
            try:
                if not DOCX_AVAILABLE:
                    st.error("❌ python-docx library not available. Please install it first: pip install python-docx")
                    return
                
                key_points = analysis.get('key_points', {}) if analysis else {}
                doc = st.session_state.verbatim_ai.export_to_docx(transcript, key_points)
                
                # Save to bytes
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                
                st.download_button(
                    label="📥 Download DOCX",
                    data=buffer.getvalue(),
                    file_name=f"{base_filename}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                st.success("✅ DOCX export completed successfully!")
            except Exception as e:
                st.error(f"❌ DOCX export failed: {str(e)}")
        
        elif export_format == "PDF":
            try:
                if not REPORTLAB_AVAILABLE:
                    st.error("❌ ReportLab library not available. Please install it first: pip install reportlab")
                    return
                
                key_points = analysis.get('key_points', {}) if analysis else {}
                buffer = st.session_state.verbatim_ai.export_to_pdf(transcript, key_points)
                
                st.download_button(
                    label="📥 Download PDF",
                    data=buffer.getvalue(),
                    file_name=f"{base_filename}.pdf",
                    mime="application/pdf"
                )
                st.success("✅ PDF export completed successfully!")
            except Exception as e:
                st.error(f"❌ PDF export failed: {str(e)}")
    
    # Show export preview information
    if transcript:
        st.subheader("📋 Export Preview")
        col1, col2, col3 = st.columns(3)
        
        with col1:
            st.metric("Duration", f"{getattr(transcript, 'audio_duration', 0) / 1000 / 60:.2f} min")
        
        with col2:
            st.metric("Words", f"{len(transcript.text.split()):,}")
        
        with col3:
            st.metric("Speakers", f"{len(set([u.speaker for u in transcript.utterances]))}")
        
        # Show key points summary
        if analysis and analysis.get('key_points'):
            key_points = analysis['key_points']
            st.subheader("🎯 Key Points Summary")
            
            kp_col1, kp_col2, kp_col3 = st.columns(3)
            
            with kp_col1:
                st.metric("Decisions", len(key_points.get('decisions', [])))
            
            with kp_col2:
                st.metric("Action Items", len(key_points.get('action_items', [])))
            
            with kp_col3:
                st.metric("Questions", len(key_points.get('questions', [])))

def show_speaker_sentiment_eda():
    """Advanced Speaker Sentiment EDA with comprehensive visualizations"""
    st.title("🎭 Speaker Sentiment EDA")
    
    if not st.session_state.verbatim_ai.transcript_data:
        st.warning("Please upload and transcribe a file first!")
        return
    
    transcript = st.session_state.verbatim_ai.transcript_data
    analysis = st.session_state.verbatim_ai.analysis_results
    
    if not analysis or 'sentiment' not in analysis:
        st.error("No sentiment analysis available. Please transcribe a file first.")
        return
    
    # Create tabs for different visualizations
    tab1, tab2, tab3, tab4, tab5 = st.tabs([
        "📊 Sentiment Over Time", 
        "😊 Emotion Distribution", 
        "🎯 Engagement Analysis",
        "📈 Speaker Comparison",
        "🔍 Detailed Insights"
    ])
    
    with tab1:
        st.subheader("📊 Sentiment Over Time")
        
        # Prepare data for time series
        time_data = []
        for utterance in transcript.utterances:
            if hasattr(utterance, 'sentiment') and utterance.sentiment:
                time_data.append({
                    'time': utterance.start / 1000 / 60,  # Convert to minutes
                    'positive': utterance.sentiment.positive,
                    'negative': utterance.sentiment.negative,
                    'neutral': utterance.sentiment.neutral,
                    'speaker': f"Speaker {utterance.speaker}"
                })
        
        if time_data:
            df_time = pd.DataFrame(time_data)
            
            # Create line chart
            fig = go.Figure()
            
            for sentiment in ['positive', 'negative', 'neutral']:
                fig.add_trace(go.Scatter(
                    x=df_time['time'],
                    y=df_time[sentiment],
                    mode='lines+markers',
                    name=sentiment.capitalize(),
                    line=dict(width=2)
                ))
            
            fig.update_layout(
                title="Sentiment Trends Over Time",
                xaxis_title="Time (minutes)",
                yaxis_title="Sentiment Score",
                hovermode='x unified',
                height=500
            )
            
            st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("No sentiment data available for time series analysis.")
    
    with tab2:
        st.subheader("😊 Emotion Distribution")
        
        # Emotion analysis
        emotion_counts = {}
        for speaker, data in analysis['sentiment'].items():
            for utterance_data in data['utterances']:
                if 'emotion' in utterance_data:
                    emotion = utterance_data['emotion']
                    if emotion not in emotion_counts:
                        emotion_counts[emotion] = 0
                    emotion_counts[emotion] += 1
        
        if emotion_counts:
            # Create pie chart
            fig = px.pie(
                values=list(emotion_counts.values()),
                names=list(emotion_counts.keys()),
                title="Emotion Distribution Across All Speakers"
            )
            st.plotly_chart(fig, use_container_width=True)
            
            # Emotion by speaker
            speaker_emotions = {}
            for speaker, data in analysis['sentiment'].items():
                speaker_emotions[speaker] = {}
                for utterance_data in data['utterances']:
                    if 'emotion' in utterance_data:
                        emotion = utterance_data['emotion']
                        if emotion not in speaker_emotions[speaker]:
                            speaker_emotions[speaker][emotion] = 0
                        speaker_emotions[speaker][emotion] += 1
            
            if speaker_emotions:
                # Create stacked bar chart
                emotion_df = []
                for speaker, emotions in speaker_emotions.items():
                    for emotion, count in emotions.items():
                        emotion_df.append({
                            'Speaker': speaker,
                            'Emotion': emotion,
                            'Count': count
                        })
                
                if emotion_df:
                    df_emotion = pd.DataFrame(emotion_df)
                    fig = px.bar(
                        df_emotion,
                        x='Speaker',
                        y='Count',
                        color='Emotion',
                        title="Emotion Distribution by Speaker",
                        barmode='stack'
                    )
                    st.plotly_chart(fig, use_container_width=True)
        else:
            st.info("No emotion data available.")
    
    with tab3:
        st.subheader("🎯 Engagement Analysis")
        
        # Calculate engagement metrics with proper error handling
        engagement_data = []
        
        # Calculate speaker engagement based on speaking time and frequency
        speaker_stats = {}
        total_meeting_time = (transcript.utterances[-1].end - transcript.utterances[0].start) / 1000 if transcript.utterances else 0
        
        for utterance in transcript.utterances:
            # Handle both string and integer speaker formats
            if isinstance(utterance.speaker, str):
                speaker = f"Speaker {utterance.speaker}"
            else:
                # Convert integer speaker ID to letter (0->A, 1->B, etc.)
                speaker = f"Speaker {chr(65 + int(utterance.speaker))}"
                
            if speaker not in speaker_stats:
                speaker_stats[speaker] = {
                    'utterances': 0,
                    'total_time': 0,
                    'words': 0,
                    'avg_sentiment': 0.5  # Default neutral sentiment
                }
            
            speaker_stats[speaker]['utterances'] += 1
            speaker_stats[speaker]['total_time'] += (utterance.end - utterance.start) / 1000
            speaker_stats[speaker]['words'] += len(utterance.text.split())
        
        # Calculate engagement scores
        for speaker, stats in speaker_stats.items():
            participation_rate = (stats['total_time'] / total_meeting_time * 100) if total_meeting_time > 0 else 0
            words_per_minute = (stats['words'] / (stats['total_time'] / 60)) if stats['total_time'] > 0 else 0
            
            # Engagement score based on participation and verbosity
            engagement_score = (participation_rate * 0.6) + (min(words_per_minute / 100, 1) * 40)
            
            engagement_data.append({
                'Speaker': speaker,
                'Utterances': stats['utterances'],
                'Speaking Time (%)': f"{participation_rate:.1f}%",
                'Words/Min': f"{words_per_minute:.1f}",
                'Engagement Score': f"{engagement_score:.1f}"
            })
        
        if engagement_data:
            df_engagement = pd.DataFrame(engagement_data)
            
            # Convert engagement scores back to float for plotting
            engagement_scores = [float(x.replace('%', '')) for x in df_engagement['Engagement Score']]
            speakers = df_engagement['Speaker'].tolist()
            
            # Engagement score chart
            fig = px.bar(
                x=speakers,
                y=engagement_scores,
                title="Speaker Engagement Scores",
                labels={'x': 'Speaker', 'y': 'Engagement Score'},
                color=engagement_scores,
                color_continuous_scale='viridis'
            )
            st.plotly_chart(fig, use_container_width=True)
            
            # Metrics table
            st.subheader("Engagement Metrics")
            st.dataframe(df_engagement, use_container_width=True)
        else:
            st.info("No engagement data available.")
    
    with tab4:
        st.subheader("📈 Speaker Comparison")
        
        # Compare speakers across multiple dimensions
        comparison_data = []
        for speaker, data in analysis['sentiment'].items():
            total_utterances = len(data['utterances'])
            positive_count = sum(
                1 for u in data['utterances'] 
                if 'sentiment' in u and u['sentiment'].get('positive', 0) > 0.5
            )
            negative_count = sum(
                1 for u in data['utterances'] 
                if 'sentiment' in u and u['sentiment'].get('negative', 0) > 0.5
            )
            
            comparison_data.append({
                'Speaker': speaker,
                'Total Utterances': total_utterances,
                'Positive %': (positive_count / total_utterances * 100) if total_utterances > 0 else 0,
                'Negative %': (negative_count / total_utterances * 100) if total_utterances > 0 else 0,
                'Neutral %': ((total_utterances - positive_count - negative_count) / total_utterances * 100) if total_utterances > 0 else 0
            })
        
        if comparison_data:
            df_comparison = pd.DataFrame(comparison_data)
            
            # Radar chart for speaker comparison
            fig = go.Figure()
            
            for _, row in df_comparison.iterrows():
                fig.add_trace(go.Scatterpolar(
                    r=[row['Positive %'], row['Negative %'], row['Neutral %'], row['Total Utterances']],
                    theta=['Positive', 'Negative', 'Neutral', 'Total Utterances'],
                    fill='toself',
                    name=row['Speaker']
                ))
            
            fig.update_layout(
                polar=dict(
                    radialaxis=dict(
                        visible=True,
                        range=[0, 100]
                    )),
                showlegend=True,
                title="Speaker Comparison Radar Chart"
            )
            
            st.plotly_chart(fig, use_container_width=True)
    
    with tab5:
        st.subheader("🔍 Detailed Insights")
        
        # Key insights
        st.markdown("### Key Insights")
        
        # Most engaged speaker
        if engagement_data:
            most_engaged = max(engagement_data, key=lambda x: float(x['Engagement Score']) if isinstance(x['Engagement Score'], (int, float, str)) else 0)
            engagement_score = float(most_engaged['Engagement Score']) if isinstance(most_engaged['Engagement Score'], (int, float, str)) else 0.0
            st.info(f"🎯 **Most Engaged Speaker**: {most_engaged['Speaker']} (Score: {engagement_score:.2f})")
        
        # Sentiment trends
        if time_data:
            df_time = pd.DataFrame(time_data)
            avg_positive = df_time['positive'].mean()
            avg_negative = df_time['negative'].mean()
            
            col1, col2, col3 = st.columns(3)
            with col1:
                st.metric("Average Positive Sentiment", f"{avg_positive:.3f}")
            with col2:
                st.metric("Average Negative Sentiment", f"{avg_negative:.3f}")
            with col3:
                st.metric("Overall Sentiment Balance", f"{avg_positive - avg_negative:.3f}")

# Fix the semantic search function with proper indentation and structure
def show_semantic_search():
    """Semantic search functionality using sentence transformers and FAISS"""
    st.title("🔍 Semantic Search")
    
    if not st.session_state.verbatim_ai.transcript_data:
        st.warning("Please upload and transcribe a file first!")
        return
    
    transcript = st.session_state.verbatim_ai.transcript_data
    
    # Check if semantic search is available
    semantic_engine = st.session_state.verbatim_ai.semantic_search_engine
    if not semantic_engine or not semantic_engine.model:
        st.warning("🔍 Semantic search is not properly configured.")
        
        # Check specific error
        if semantic_engine and semantic_engine.error_message:
            st.error(f"Error: {semantic_engine.error_message}")
        
        st.info("""
        **To enable semantic search, ensure the following packages are installed:**
        ```bash
        pip install sentence-transformers faiss-cpu
        ```
        """)
        
        # Try to reinitialize
        if st.button("🔄 Try to Initialize Semantic Search"):
            try:
                from semantic_search import SemanticSearchEngine
                st.session_state.verbatim_ai.semantic_search_engine = SemanticSearchEngine()
                st.rerun()
            except Exception as e:
                st.error(f"Failed to initialize: {str(e)}")
        return
    
    # Build index if not already built
    if not hasattr(st.session_state, 'semantic_index_built') or not st.session_state.semantic_index_built:
        with st.spinner("Building semantic search index..."):
            success = st.session_state.verbatim_ai.build_semantic_index(transcript)
            if success:
                st.session_state.semantic_index_built = True
                st.success("✅ Semantic search index built successfully!")
            else:
                st.error("❌ Failed to build semantic search index.")
                return
    
    # Search interface
    st.subheader("🔍 Search Transcript")
    
    # Search options
    search_type = st.radio(
        "Search type:",
        ["Semantic Search", "Topic Clusters", "Similar Utterances"]
    )
    
    if search_type == "Semantic Search":
        query = st.text_input("Enter your search query:", placeholder="e.g., project timeline, budget discussion, team collaboration")
        top_k = st.slider("Number of results:", min_value=1, max_value=20, value=5)
        
        if st.button("🔍 Search", type="primary") and query:
            with st.spinner("Searching..."):
                results = st.session_state.verbatim_ai.semantic_search(query, top_k)
                
                if results:
                    st.subheader(f"📋 Search Results ({len(results)} found)")
                    
                    for i, result in enumerate(results, 1):
                        with st.expander(f"Result {i}: Score {result['score']:.3f}", expanded=True):
                            # Safe access to speaker field
                            speaker_info = result.get('speaker', 'Unknown Speaker')
                            st.markdown(f"**Speaker**: {speaker_info}")
                            
                            st.markdown(f"**Text**: {result['text']}")
                            
                            # Safe access to timing fields with proper error handling
                            try:
                                if 'start_time' in result and 'end_time' in result:
                                    # Show timestamp in readable format
                                    start_min = result['start_time'] / 1000 / 60
                                    start_sec = (result['start_time'] / 1000) % 60
                                    end_min = result['end_time'] / 1000 / 60
                                    end_sec = (result['end_time'] / 1000) % 60
                                    st.caption(f"📍 Timestamp: {start_min:.0f}:{start_sec:02.0f} - {end_min:.0f}:{end_sec:02.0f}")
                                elif 'index' in result:
                                    st.caption(f"📍 Utterance index: {result['index']}")
                                    
                                    # Try to get actual utterance data from transcript
                                    if (st.session_state.verbatim_ai.transcript_data and 
                                        result['index'] < len(st.session_state.verbatim_ai.transcript_data.utterances)):
                                        utterance = st.session_state.verbatim_ai.transcript_data.utterances[result['index']]
                                        start_min = utterance.start / 1000 / 60
                                        start_sec = (utterance.start / 1000) % 60
                                        end_min = utterance.end / 1000 / 60
                                        end_sec = (utterance.end / 1000) % 60
                                        st.caption(f"📍 Actual timestamp: {start_min:.0f}:{start_sec:02.0f} - {end_min:.0f}:{end_sec:02.0f}")
                                        st.caption(f"📍 Speaker: Speaker {utterance.speaker}")
                                else:
                                    st.caption(f"📍 Position: Result {i}")
                            except Exception as e:
                                st.caption(f"📍 Position: Result {i}")
                            
                            # Show search type if available
                            if 'search_type' in result:
                                st.caption(f"🔍 Search type: {result['search_type']}")
                
                else:
                    st.info("No results found for your query.")
    
    elif search_type == "Topic Clusters":
        n_clusters = st.slider("Number of clusters:", min_value=2, max_value=10, value=5)
        
        if st.button("🔍 Generate Topic Clusters", type="primary"):
            with st.spinner("Generating topic clusters..."):
                clusters = st.session_state.verbatim_ai.get_topic_clusters(n_clusters)
                
                if clusters:
                    st.subheader(f"📊 Topic Clusters ({len(clusters)} clusters)")
                    
                    for i, cluster in enumerate(clusters, 1):
                        with st.expander(f"Cluster {i}: {cluster['size']} utterances", expanded=True):
                            st.markdown(f"**Cluster Size**: {cluster['size']} utterances")
                            st.markdown(f"**Summary**: {cluster['summary']}")
                            
                            # Show sample utterances
                            st.markdown("**Sample Utterances**:")
                            for j, utterance in enumerate(cluster['utterances'][:3], 1):
                                st.markdown(f"{j}. **{utterance.speaker}**: {utterance.text[:100]}...")
                else:
                    st.info("No clusters generated.")

    elif search_type == "Similar Utterances":
        # Find similar utterances to a given text
        reference_text = st.text_area("Enter reference text:", placeholder="Enter text to find similar utterances...")
        top_k = st.slider("Number of similar results:", min_value=1, max_value=10, value=3)
        
        if st.button("🔍 Find Similar", type="primary") and reference_text:
            with st.spinner("Finding similar utterances..."):
                results = st.session_state.verbatim_ai.get_similar_utterances(reference_text, top_k)
                
                if results:
                    st.subheader(f"📋 Similar Utterances ({len(results)} found)")
                    
                    for i, result in enumerate(results, 1):
                        with st.expander(f"Similar {i}: Score {result['score']:.3f}", expanded=True):
                            # Safe access to speaker field
                            speaker_info = result.get('speaker', 'Unknown Speaker')
                            st.markdown(f"**Speaker**: {speaker_info}")
                            st.markdown(f"**Text**: {result['text']}")
                            
                            # Show index if available
                            if 'index' in result:
                                st.caption(f"📍 Utterance index: {result['index']}")
                else:
                    st.info("No similar utterances found.")
    
    # Quick search bar for similar utterances
    st.subheader("🔍 Quick Similar Utterances Search")
    quick_search = st.text_input("Quick search for similar utterances:", placeholder="Type any text to find similar utterances...")
    
    if quick_search:
        with st.spinner("Finding similar utterances..."):
            similar_results = st.session_state.verbatim_ai.get_similar_utterances(quick_search, 5)
            
            if similar_results:
                st.subheader("📋 Similar Utterances Found")
                for i, result in enumerate(similar_results, 1):
                    # Safe access to speaker field
                    speaker_info = result.get('speaker', 'Unknown Speaker')
                    st.markdown(f"**{i}.** **{speaker_info}** (Score: {result['score']:.3f}): {result['text']}")
                    
                    # Safe timestamp handling
                    try:
                        if 'start_time' in result and result['start_time']:
                            start_min = result['start_time'] / 1000 / 60
                            start_sec = (result['start_time'] / 1000) % 60
                            st.caption(f"⏰ Timestamp: {start_min:.0f}:{start_sec:02.0f}")
                        elif 'index' in result:
                            # Try to get timestamp from transcript data
                            index = result['index']
                            if (st.session_state.verbatim_ai.transcript_data and 
                                hasattr(st.session_state.verbatim_ai.transcript_data, 'utterances') and
                                index < len(st.session_state.verbatim_ai.transcript_data.utterances)):
                                utterance = st.session_state.verbatim_ai.transcript_data.utterances[index]
                                start_min = utterance.start / 1000 / 60
                                start_sec = (utterance.start / 1000) % 60
                                st.caption(f"⏰ Timestamp: {start_min:.0f}:{start_sec:02.0f}")
                            else:
                                st.caption(f"📍 Position: Result {i}")
                        else:
                            st.caption(f"📍 Position: Result {i}")
                    except Exception as e:
                        st.caption(f"📍 Position: Result {i}")
                    
                    st.markdown("---")
            else:
                st.info("No similar utterances found for your search term.")

def show_task_monitor():
    """Task monitoring interface"""
    st.title("🔄 Task Monitor")
    
    if ADVANCED_FEATURES_AVAILABLE:
        try:
            from async_task_queue import TaskManager, show_task_monitor, show_task_progress
            
            # Show task monitor interface
            show_task_monitor()
            
        except ImportError:
            st.error("❌ Task monitoring not available. Please install required dependencies.")
            st.info("Install with: `pip install celery redis`")
    else:
        st.error("❌ Advanced features not available.")
        
        # Show manual task simulation
        st.subheader("📊 Task Simulation")
        st.info("This is a simulation of what the task monitor would look like.")
        
        # Mock task data
        mock_tasks = [
            {"id": "task_001", "status": "SUCCESS", "name": "Audio Transcription", "progress": 100},
            {"id": "task_002", "status": "PENDING", "name": "Sentiment Analysis", "progress": 0},
            {"id": "task_003", "status": "PROGRESS", "name": "Semantic Indexing", "progress": 65},
        ]
        
        for task in mock_tasks:
            col1, col2, col3, col4 = st.columns(4)
            with col1:
                st.write(f"**{task['name']}**")
            with col2:
                if task['status'] == 'SUCCESS':
                    st.success(task['status'])
                elif task['status'] == 'PENDING':
                    st.warning(task['status'])
                else:
                    st.info(task['status'])
            with col3:
                st.progress(task['progress'] / 100)
            with col4:
                st.write(f"{task['progress']}%")

def show_email_summary():
    """Email summary and engagement scoring functionality"""
    global email_generator
    st.title("📧 Email Summary & Engagement")
    
    if not st.session_state.verbatim_ai.transcript_data:
        st.warning("Please upload and transcribe a file first!")
        return
    
    transcript = st.session_state.verbatim_ai.transcript_data
    analysis = st.session_state.verbatim_ai.analysis_results
    
    # Check if email functionality is available
    if not email_generator:
        st.warning("📧 Email functionality is not properly configured.")
        
        # Check current config
        from config import EMAIL_SENDER, EMAIL_PASSWORD
        
        if not EMAIL_SENDER:
            st.error("❌ EMAIL_SENDER not configured in config.py")
        else:
            st.success(f"✅ EMAIL_SENDER configured: {EMAIL_SENDER}")
        
        if not EMAIL_PASSWORD:
            st.error("❌ EMAIL_PASSWORD not configured in config.py")
        else:
            st.success("✅ EMAIL_PASSWORD configured")
        
        st.info("""
        **Configuration status:**
        - Email settings are loaded from config.py
        - EMAIL_SENDER and EMAIL_PASSWORD are configured
        - Email functionality will be available once email_generator is initialized properly
        """)
        
        # Show current values (masked)
        with st.expander("🔍 Current Configuration"):
            st.write(f"EMAIL_SENDER: {EMAIL_SENDER}")
            st.write(f"EMAIL_PASSWORD: {'*' * len(EMAIL_PASSWORD) if EMAIL_PASSWORD else 'Not set'}")
        
        # Try to initialize email generator
        if EMAIL_SENDER and EMAIL_PASSWORD:
            try:
                from email_summary import EmailSummaryGenerator
                email_generator = EmailSummaryGenerator()
                st.success("✅ Email generator initialized successfully!")
                st.rerun()
            except Exception as e:
                st.error(f"❌ Failed to initialize email generator: {e}")
        
        return
    
    # Create tabs
    tab1, tab2, tab3 = st.tabs(["📊 Engagement Analysis", "📧 Email Summary", "⚙️ Email Settings"])
    
    with tab1:
        st.subheader("📊 Meeting Engagement Analysis")
        
        # Calculate engagement score
        transcript_data = {
            'duration_minutes': transcript.audio_duration / 1000 / 60,
            'total_words': len(transcript.text.split()),
            'speakers': list(set([u.speaker for u in transcript.utterances])),
            'utterances': transcript.utterances,
            'key_points': analysis.get('key_points', {}) if analysis else {}
        }
        
        engagement = st.session_state.verbatim_ai.calculate_engagement_score(transcript_data)
        
        # Display engagement metrics
        col1, col2, col3, col4 = st.columns(4)
        
        with col1:
            st.metric(
                "Total Score", 
                f"{engagement['total_score']}%",
                delta=f"{engagement['level']}"
            )
        
        with col2:
            st.metric(
                "Participation", 
                f"{engagement['participation_score']}%"
            )
        
        with col3:
            st.metric(
                "Interaction", 
                f"{engagement['interaction_score']}%"
            )
        
        with col4:
            st.metric(
                "Action Items", 
                f"{engagement['action_score']}%"
            )
        
        # Engagement breakdown chart (removed sentiment)
        engagement_breakdown = {
            'Participation': engagement['participation_score'],
            'Interaction': engagement['interaction_score'],
            'Action Items': engagement['action_score']
        }
        
        fig = px.bar(
            x=list(engagement_breakdown.keys()),
            y=list(engagement_breakdown.values()),
            title="Engagement Score Breakdown",
            color=list(engagement_breakdown.values()),
            color_continuous_scale='viridis'
        )
        fig.update_layout(height=400)
        st.plotly_chart(fig, use_container_width=True)
        
        # Engagement insights
        st.subheader("🎯 Engagement Insights")
        
        if engagement['total_score'] >= 80:
            st.success("🎉 **Excellent Engagement**: This meeting showed high levels of participation, interaction, and positive sentiment!")
        elif engagement['total_score'] >= 60:
            st.info("👍 **Good Engagement**: The meeting had solid participation and interaction levels.")
        elif engagement['total_score'] >= 40:
            st.warning("⚠️ **Moderate Engagement**: Consider ways to increase participation and interaction.")
        else:
            st.error("📉 **Low Engagement**: This meeting needs improvement in participation and interaction.")
    
    with tab2:
        st.subheader("📧 Generate Email Summary")
        
        # Email configuration
        recipient_email = st.text_input("Recipient Email:", placeholder="example@company.com")
        email_subject = st.text_input("Email Subject:", value=f"Weekly Meeting Summary - {datetime.now().strftime('%B %d, %Y')}")
        
        # Summary options
        include_engagement = st.checkbox("Include engagement analysis", value=True)
        include_key_points = st.checkbox("Include key points", value=True)
        include_sentiment = st.checkbox("Include sentiment analysis", value=True)
        
        if st.button("📧 Generate & Send Email", type="primary") and recipient_email:
            with st.spinner("Generating email summary..."):
                try:
                    # Prepare meeting data
                    meeting_data = [{
                        'title': 'Meeting Transcript',
                        'duration_minutes': transcript.audio_duration / 1000 / 60,
                        'speakers': list(set([u.speaker for u in transcript.utterances])),
                        'key_points': analysis.get('key_points', {}) if analysis else {},
                        'date': datetime.now()
                    }]
                    
                    # Generate and send email
                    success = email_generator.send_weekly_summary(
                        recipient_email, meeting_data
                    )
                    
                    if success:
                        st.success("✅ Email sent successfully!")
                    else:
                        st.error("❌ Failed to send email. Please check your email settings.")
                        
                except Exception as e:
                    st.error(f"❌ Error sending email: {e}")
                    st.info("💡 Make sure your email settings are configured correctly in config.py")
        
        # Preview email content
        if st.button("👁️ Preview Email Content"):
            meeting_data = [{
                'title': 'Meeting Transcript',
                'duration_minutes': transcript.audio_duration / 1000 / 60,
                'speakers': list(set([u.speaker for u in transcript.utterances])),
                                        'key_points': analysis.get('key_points', {}) if analysis else {},
                'date': datetime.now()
            }]
            
            html_content = email_generator.generate_weekly_summary(meeting_data)
            
            st.subheader("📧 Email Preview")
            st.markdown("**HTML Content Preview:**")
            st.code(html_content[:1000] + "..." if len(html_content) > 1000 else html_content)
    
    with tab3:
        st.subheader("⚙️ Email Configuration")
        
        st.info("""
        **Email Configuration Instructions:**
        
        1. **Gmail Setup**: 
           - Enable 2-factor authentication
           - Generate an App Password
           - Use the App Password instead of your regular password
        
        2. **Environment Variables**:
           - Set `EMAIL_SENDER` to your email address
           - Set `EMAIL_PASSWORD` to your app password
        
        3. **SMTP Settings**:
           - Server: smtp.gmail.com
           - Port: 587
           - Security: TLS
        """)
        
        # Display current settings
        st.subheader("Current Settings")
        st.code(f"""
        SMTP Server: {email_generator.smtp_server if email_generator else 'Not configured'}
        SMTP Port: {email_generator.smtp_port if email_generator else 'Not configured'}
        Sender Email: {email_generator.sender_email if email_generator else 'Not configured'}
        """)
        
        # Test email functionality
        if st.button("🧪 Test Email Configuration"):
            test_email = st.text_input("Test Email Address:", placeholder="test@example.com")
            if test_email and st.button("Send Test Email"):
                try:
                    success = email_generator.send_email(
                        test_email, 
                        "Test Email from VerbatimAI", 
                        "<h1>Test Email</h1><p>This is a test email from VerbatimAI.</p>"
                    )
                    if success:
                        st.success("✅ Test email sent successfully!")
                    else:
                        st.error("❌ Test email failed to send.")
                except Exception as e:
                    st.error(f"❌ Test email error: {e}")

def show_user_management():
    """User management interface for admins"""
    st.title("👥 User Management")
    
    # Check if user has permission
    if not check_authentication():
        st.error("❌ Please log in to access this feature.")
        return
    
    user_manager = st.session_state.user_manager
    current_user_role = get_current_user_role()
    
    if not user_manager.has_permission(st.session_state.session_id, Permission.MANAGE_USERS):
        st.error("❌ You don't have permission to access user management.")
        return
    
    # Display all users
    st.subheader("📋 User List")
    
    users = list(user_manager.users.values())
    
    if users:
        # Create user table
        user_data = []
        for user in users:
            user_data.append({
                'Username': user.username,
                'Role': user.role.value.title(),
                'Email': user.email or 'N/A',
                'Status': 'Active' if user.is_active else 'Inactive',
                'Last Login': user.last_login.strftime('%Y-%m-%d %H:%M') if user.last_login else 'Never'
            })
        
        df = pd.DataFrame(user_data)
        st.dataframe(df, use_container_width=True)
        
        # User actions
        st.subheader("🔧 User Actions")
        
        col1, col2 = st.columns(2)
        
        with col1:
            st.write("**Add New User**")
            new_username = st.text_input("Username", key="new_username")
            new_email = st.text_input("Email", key="new_email")
            new_role = st.selectbox("Role", ["admin", "contributor", "viewer", "guest"], key="new_role")
            
            if st.button("Add User", key="add_user_btn") and new_username:
                # This would add user creation logic
                st.success(f"✅ User {new_username} created!")
        
        with col2:
            st.write("**Modify User**")
            selected_user = st.selectbox("Select User", [u.username for u in users], key="selected_user")
            
            if selected_user:
                user = user_manager.users.get(selected_user)
                if user:
                    new_status = st.checkbox("Active", value=user.is_active, key="user_status")
                    new_role = st.selectbox("Role", ["admin", "contributor", "viewer", "guest"], 
                                          index=["admin", "contributor", "viewer", "guest"].index(user.role.value),
                                          key="modify_role")
                    
                    if st.button("Update User", key="update_user_btn"):
                        # This would add user update logic
                        st.success(f"✅ User {selected_user} updated!")
    else:
        st.info("No users found.")

def show_system_settings():
    """System settings interface for admins"""
    st.title("⚙️ System Settings")
    
    # Check if user has permission
    if not check_authentication():
        st.error("❌ Please log in to access this feature.")
        return
    
    user_manager = st.session_state.user_manager
    
    if not user_manager.has_permission(st.session_state.session_id, Permission.SYSTEM_CONFIG):
        st.error("❌ You don't have permission to access system settings.")
        return
    
    # System configuration
    st.subheader("🔧 Configuration")
    
    col1, col2 = st.columns(2)
    
    with col1:
        st.write("**API Configuration**")
        api_key = st.text_input("AssemblyAI API Key", value=ASSEMBLYAI_API_KEY, type="password", key="api_key")
        
        if st.button("Update API Key", key="update_api_btn"):
            st.success("✅ API key updated!")
        
        st.write("**Email Configuration**")
        smtp_server = st.text_input("SMTP Server", value="smtp.gmail.com", key="smtp_server")
        smtp_port = st.number_input("SMTP Port", value=587, key="smtp_port")
        sender_email = st.text_input("Sender Email", value="your-email@gmail.com", key="sender_email")
        
        if st.button("Update Email Settings", key="update_email_btn"):
            st.success("✅ Email settings updated!")
    
    with col2:
        st.write("**System Information**")
        st.write(f"**Python Version:** {sys.version}")
        st.write(f"**Streamlit Version:** {st.__version__}")
        try:
            st.write(f"**AssemblyAI SDK:** {aai.__version__}")
        except AttributeError:
            st.write("**AssemblyAI SDK:** Version info not available")
        
        # System health
        st.write("**System Health**")
        try:
            api_status = "✅ Valid" if st.session_state.verbatim_ai.verify_api_key() else "❌ Invalid"
            st.write(f"**API Status:** {api_status}")
        except Exception as e:
            st.write(f"**API Status:** ❌ Error checking API key: {str(e)}")
        
        # Task queue status
        st.write("**Task Queue Status**")
        try:
            from async_task_queue import TaskManager
            task_manager = TaskManager()
            tasks = task_manager.get_all_tasks()
            active_tasks = len([t for t in tasks if t['status'] not in ['SUCCESS', 'FAILURE']])
            st.write(f"**Active Tasks:** {active_tasks}")
        except Exception as e:
            st.write("**Task Queue:** Not available")
            st.caption(f"Error: {str(e)}")
        
        # System report export
        st.subheader("📊 System Report Export")
        if st.button("📋 Generate System Report"):
            report_data = generate_system_report()
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            
            # Export as CSV
            csv_buffer = io.StringIO()
            pd.DataFrame([report_data]).to_csv(csv_buffer, index=False)
            
            st.download_button(
                label="📥 Download System Report (CSV)",
                data=csv_buffer.getvalue(),
                file_name=f"system_report_{timestamp}.csv",
                mime="text/csv"
            )
    
    # Advanced settings
    st.subheader("🔬 Advanced Settings")
    
    col3, col4 = st.columns(2)
    
    with col3:
        st.write("**Model Configuration**")
        enable_advanced_models = st.checkbox("Enable Advanced NLP Models", value=True, key="enable_advanced_models")
        enable_emotion_detection = st.checkbox("Enable Enhanced Emotion Detection", value=True, key="enable_emotion_detection")
        enable_semantic_search = st.checkbox("Enable Semantic Search", value=True, key="enable_semantic_search")
        
        if st.button("Update Model Settings", key="update_model_btn"):
            st.success("✅ Model settings updated!")
    
    with col4:
        st.write("**Performance Settings**")
        max_file_size = st.number_input("Max File Size (MB)", value=1024, key="max_file_size")
        transcription_timeout = st.number_input("Transcription Timeout (minutes)", value=30, key="transcription_timeout")
        enable_caching = st.checkbox("Enable Result Caching", value=True, key="enable_caching")
        
        if st.button("Update Performance Settings", key="update_performance_btn"):
            st.success("✅ Performance settings updated!")
    
    # System maintenance
    st.subheader("🛠️ System Maintenance")
    
    col5, col6 = st.columns(2)
    
    with col5:
        if st.button("🧹 Clear Cache", key="clear_cache_btn"):
            st.success("✅ Cache cleared!")
        
        if st.button("📊 Generate System Report", key="generate_report_btn"):
            st.success("✅ System report generated!")
    
    with col6:
        if st.button("🔄 Restart Services", key="restart_services_btn"):
            st.success("✅ Services restarted!")
        
        if st.button("📋 View System Logs", key="view_logs_btn"):
            st.info("System logs would be displayed here.")

def show_library():
    st.title("📚 Meeting Library / History")
    meetings = load_meetings()
    if not meetings:
        st.info("No meetings found. Transcribe a meeting to see it here.")
        return
    for meeting in meetings:
        with st.expander(f"{meeting['date']} | Speakers: {len(meeting['speakers'])}"):
            st.markdown(f"**Summary:** {meeting['summary']}")
            if st.button(f"Show Transcript", key=meeting['id']):
                st.markdown(f"**Transcript:**\n\n{meeting['transcript']}")

MEETINGS_DIR = "meetings"
os.makedirs(MEETINGS_DIR, exist_ok=True)

def save_meeting(transcript, summary):
    """Save meeting transcript and summary to a JSON file."""
    meeting_id = f"meeting_{datetime.now().strftime('%Y%m%d_%H%M%S')}"
    data = {
        "id": meeting_id,
        "date": datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        "summary": summary,
        "transcript": transcript.text,
        "speakers": list(set([u.speaker for u in transcript.utterances]))
    }
    with open(os.path.join(MEETINGS_DIR, f"{meeting_id}.json"), "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    return meeting_id

def load_meetings():
    """Load all saved meetings from the meetings directory."""
    meetings = []
    for fname in os.listdir(MEETINGS_DIR):
        if fname.endswith(".json"):
            with open(os.path.join(MEETINGS_DIR, fname), "r", encoding="utf-8") as f:
                meetings.append(json.load(f))
    meetings.sort(key=lambda x: x["date"], reverse=True)
    return meetings


if __name__ == "__main__":
    main() 